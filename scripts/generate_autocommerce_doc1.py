from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
ASSETS = Path("/tmp/autocommerce-doc1-assets")
OUTPUT = ROOT / "artifacts" / "Autocommerce" / "document-01"

CHARCOAL = "202020"
RED = "E00810"
GREEN = "104830"
LIGHT = "F4F5F6"
MID = "D9DDE1"
WHITE = "FFFFFF"

SIGNATURE = (
    "PREPARED BY",
    "DJIGO DJIBI",
    "CX Consultant | Strategic Product & UX/UI Designer | AI Front-End Developer",
)

FR = {
    "code": "FR",
    "title": "UX/UI Audit Summary\nAnalyse fonctionnelle approfondie",
    "subtitle": "AUTOCOMMERCE · DOCUMENT 01 · AUDIT UX/UI",
    "toc": "Table des matières",
    "confidential": "AUTOCOMMERCE · CONFIDENTIEL",
    "pages": [
        (
            "Introduction et mandat",
            [
                (
                    "Objectif",
                    "Évaluer les écrans existants d’AutoCommerce, transformer les annotations "
                    "fournies en exigences actionnables et définir une architecture fonctionnelle "
                    "pour un MVP Figma/Make distinctif, simple et crédible.",
                ),
                (
                    "Positionnement",
                    "AutoCommerce demeure propriétaire du produit. La plateforme sert trois "
                    "publics : acheteurs, vendeurs/commerçants et équipe d’administration "
                    "AutoCommerce. Cars.ca constitue la référence de couverture fonctionnelle, "
                    "sans reproduction de son identité ou de son interface.",
                ),
                (
                    "Principe directeur",
                    "Simple au premier regard, puissant à la demande. Les véhicules et l’action "
                    "principale restent dominants; filtres, analyses et fonctions expertes se "
                    "déploient progressivement.",
                ),
            ],
        ),
        (
            "Périmètre observé et méthode",
            [
                (
                    "Corpus",
                    "Le corpus comprend 18 pages de captures annotées : accueil, navigation Buy, "
                    "inventaire, modes liste et galerie, fiche véhicule, achat immédiat, enchères, "
                    "vente, authentification, tableau de bord et assistant Add Vehicle.",
                ),
                (
                    "Méthode d’évaluation",
                    "Chaque écran est analysé selon l’objectif utilisateur, la hiérarchie visuelle, "
                    "la charge cognitive, la cohérence, les erreurs possibles, l’accessibilité, "
                    "la confiance et la capacité responsive. Les constats sont rapprochés des "
                    "heuristiques de Nielsen Norman Group et de WCAG 2.2 AA.",
                ),
                (
                    "Limites",
                    "Les captures ne démontrent pas tous les comportements, règles métier ni états "
                    "d’erreur. Les recommandations décrivent donc les exigences à valider dans le "
                    "prototype, et non un audit du code de production.",
                ),
            ],
        ),
        (
            "Accueil et navigation globale",
            [
                (
                    "Constat",
                    "L’accueil repose sur une forte dominante noire et une animation automobile. "
                    "L’identité est reconnaissable, mais la proposition de valeur, les chemins "
                    "Acheter/Vendre et la prochaine action ne sont pas suffisamment explicites.",
                ),
                (
                    "Recommandation",
                    "Conserver le logo, le noir, le rouge et le vert historiques, tout en faisant "
                    "du blanc et du gris clair les surfaces principales. Installer une recherche "
                    "centrale, deux entrées Acheter/Vendre, des preuves de confiance et un accès "
                    "direct au tableau de bord.",
                ),
                (
                    "Navigation cible",
                    "En-tête compact et persistant : logo, Acheter, Vendre, Enchères, Services, "
                    "Aide, langue et compte. Sur mobile, conserver Recherche, Favoris, Vendre et "
                    "Compte dans une navigation basse accessible au pouce.",
                ),
            ],
        ),
        (
            "Inventaire, recherche et filtres",
            [
                (
                    "Problème principal",
                    "L’annotation confirme que le filtre latéral statique prend trop de place. "
                    "Il réduit la surface d’inventaire et concurrence les véhicules. Le méga-menu "
                    "Inventory occupe également l’écran avec trop d’options simultanées.",
                ),
                (
                    "Solution fonctionnelle",
                    "Remplacer le panneau permanent par un tiroir rétractable sur ordinateur et "
                    "une feuille modale sur mobile. Les critères actifs deviennent des pastilles "
                    "supprimables. La fermeture du filtre conserve les choix et restaure toute la "
                    "largeur des résultats.",
                ),
                (
                    "Filtres prioritaires",
                    "Localisation, distance, marque, modèle, année, prix, kilométrage, carrosserie, "
                    "carburant, transmission, traction, état, vendeur et disponibilité. Les filtres "
                    "experts restent regroupés dans « Plus de critères ».",
                ),
            ],
        ),
        (
            "Résultats et fiche véhicule",
            [
                (
                    "Modes d’exploration",
                    "Les modes liste et galerie existent déjà. Ils doivent partager les mêmes "
                    "données, filtres et états. Une vue carte peut être ajoutée ultérieurement. "
                    "Le changement de vue ne doit jamais réinitialiser la recherche.",
                ),
                (
                    "Carte véhicule",
                    "Photo principale, année-marque-modèle, version, kilométrage, localisation, "
                    "prix total, badge de confiance, vendeur et actions Favori/Comparer. Les "
                    "actions Acheter maintenant, Faire une offre et Contacter doivent respecter "
                    "une hiérarchie constante.",
                ),
                (
                    "Fiche détaillée",
                    "Galerie plein écran, résumé du prix, historique, état, équipements, vendeur, "
                    "financement indicatif, disponibilité et actions. Les informations critiques "
                    "restent visibles pendant le défilement sans masquer le contenu.",
                ),
            ],
        ),
        (
            "Achat immédiat et enchères",
            [
                (
                    "Achat immédiat",
                    "Le parcours Buy Now doit afficher le prix total disponible, les frais "
                    "obligatoires, les conditions, l’identité du vendeur et la prochaine étape "
                    "avant tout engagement. Aucun prix d’appel inaccessible ne doit être utilisé.",
                ),
                (
                    "Offre",
                    "Faire une offre nécessite un montant, une échéance, un récapitulatif et une "
                    "confirmation. Les états envoyée, consultée, acceptée, refusée, expirée et "
                    "contre-offre doivent être visibles dans le tableau de bord.",
                ),
                (
                    "Enchères",
                    "L’état vide actuel doit devenir utile : explication, calendrier, alertes et "
                    "véhicules à venir. Une enchère active exige heure serveur, mise actuelle, "
                    "incrément, dépôt, historique, prolongation anti-sniping et confirmation.",
                ),
            ],
        ),
        (
            "Vente et cycle de vie des annonces",
            [
                (
                    "Cycle cible",
                    "Brouillon → contrôle qualité → en attente de validation → active → réservée "
                    "→ vendue, expirée ou refusée. Chaque statut doit expliquer la cause, les "
                    "actions disponibles et la prochaine étape.",
                ),
                (
                    "Tableau de bord commerçant",
                    "Vue synthétique des stocks, annonces à corriger, prospects chauds, rendez-vous, "
                    "offres, délais de réponse et ventes. Des vues Inventaire, Prospects, Messages, "
                    "Analytique, Équipe et Paramètres complètent le pilotage.",
                ),
                (
                    "Productivité",
                    "Recherche globale, filtres sauvegardés, actions groupées, duplication contrôlée, "
                    "import CSV/API, attribution des prospects et journal d’activité réduisent la "
                    "charge opérationnelle des commerçants.",
                ),
            ],
        ),
        (
            "Assistant d’ajout d’un véhicule",
            [
                (
                    "Structure existante",
                    "Les captures montrent quatre étapes : détails du véhicule, pneus, photos et "
                    "dommages, description. Le formulaire contient de nombreux choix utiles, mais "
                    "la densité et les listes de pastilles augmentent l’effort et le risque d’erreur.",
                ),
                (
                    "Parcours recommandé",
                    "1) VIN ou saisie manuelle; 2) caractéristiques vérifiées; 3) état et pneus; "
                    "4) photos guidées; 5) prix et description; 6) prévisualisation. Une sauvegarde "
                    "automatique et un indicateur de progression accompagnent chaque étape.",
                ),
                (
                    "Assistance intelligente",
                    "Le VIN préremplit les données disponibles sans verrouiller la correction. "
                    "L’analyse photo détecte flou, doublons et angles manquants. Une description "
                    "assistée résume les faits validés et reste entièrement modifiable.",
                ),
            ],
        ),
        (
            "Rôles, tableaux de bord et permissions",
            [
                (
                    "Acheteur",
                    "Favoris, comparateur, recherches et alertes, messages, rendez-vous, offres, "
                    "documents et progression d’achat.",
                ),
                (
                    "Vendeur ou commerçant",
                    "Inventaire, création d’annonce, score qualité, recommandations de prix, "
                    "prospects, tâches, performance, équipe et exports.",
                ),
                (
                    "AutoCommerce",
                    "Comptes, commerçants, modération, règles de publication, abonnements, fraude, "
                    "litiges, qualité, revenus, conversion, journal d’audit et contrôle d’accès.",
                ),
                (
                    "Permissions",
                    "Administrateur propriétaire, gestionnaire, vendeur, marketing et lecture seule. "
                    "Appliquer le moindre privilège, l’authentification renforcée pour les actions "
                    "sensibles et un historique traçable.",
                ),
            ],
        ),
        (
            "Architecture fonctionnelle cible",
            [
                (
                    "Couche expérience",
                    "Web responsive et prototype Figma couvrant accueil, recherche, résultats, fiche, "
                    "compte, vente et administration. Les composants partagent tokens, états et règles "
                    "d’accessibilité.",
                ),
                (
                    "Services métier",
                    "Catalogue, recherche, comptes, organisations, annonces, médias, offres, enchères, "
                    "prospects, messagerie, notifications, tarification, modération et analytique.",
                ),
                (
                    "Automatisation MVP",
                    "Make simule la réception VIN, le contrôle photo, la génération de brouillon, "
                    "l’alerte de validation et la notification de prospect. Les scénarios utilisent "
                    "des données fictives et des webhooks idempotents.",
                ),
                (
                    "Données et sécurité",
                    "Séparer utilisateurs, organisations, véhicules, annonces, médias et événements. "
                    "Consentement, minimisation, chiffrement, rétention, accès et suppression doivent "
                    "être définis avant toute production.",
                ),
            ],
        ),
        (
            "Priorités, critères d’acceptation et synthèse",
            [
                (
                    "Priorité MVP",
                    "Recherche et filtres rétractables; liste/galerie; fiche véhicule; connexion; "
                    "tableau de bord commerçant; assistant VIN/photos/description; modération "
                    "AutoCommerce; prototype responsive et accessible.",
                ),
                (
                    "Critères clés",
                    "Le filtre se ferme sans perdre les choix; toutes les actions sont utilisables "
                    "au clavier; le focus reste visible; les cibles tactiles respectent WCAG 2.2; "
                    "les erreurs sont explicites; la saisie est sauvegardée; les prix obligatoires "
                    "sont transparents; les rôles limitent réellement l’accès.",
                ),
                (
                    "Décision",
                    "AutoCommerce ne doit pas être une copie supplémentaire d’un marché automobile. "
                    "Sa différence est un système de vente assistée : inventaire plus lisible pour "
                    "l’acheteur, publication accélérée pour le commerçant et contrôle central de la "
                    "qualité pour AutoCommerce.",
                ),
            ],
        ),
    ],
    "sources": [
        "Capture AutoCommerce fournie — 18 pages annotées, 2026.",
        "AutoCommerce — routes publiques observées le 26 juillet 2026 : https://autocommerce.ca/.",
        "Nielsen Norman Group — 10 Usability Heuristics: https://www.nngroup.com/articles/ten-usability-heuristics/.",
        "Baymard Institute — Product Lists & Filtering UX: https://baymard.com/research/ecommerce-product-lists.",
        "Baymard Institute — Applied Filters Overview: https://baymard.com/blog/how-to-design-applied-filters.",
        "W3C Web Accessibility Initiative — WCAG 2.2: https://www.w3.org/TR/WCAG22/.",
        "Office de la protection de la vie privée du Canada — Consentement significatif: https://www.priv.gc.ca/fr/sujets-lies-a-la-protection-de-la-vie-privee/lois-sur-la-protection-des-renseignements-personnels-au-canada/la-loi-sur-la-protection-des-renseignements-personnels-et-les-documents-electroniques-lprpde/p_principle/principes/p_consent/.",
        "Bureau de la concurrence Canada — Indication de prix partiel: https://bureau-concurrence.canada.ca/pratiques-commerciales-trompeuses/indication-prix-partiel.",
        "Office de la consommation — Acheter ou louer un véhicule: https://ised-isde.canada.ca/site/bureau-consommation/fr/acheter-louer-articles-gros-prix/acheter-louer-vehicule.",
    ],
}

EN = {
    "code": "EN",
    "title": "UX/UI Audit Summary\nIn-depth Functional Analysis",
    "subtitle": "AUTOCOMMERCE · DOCUMENT 01 · UX/UI AUDIT",
    "toc": "Table of Contents",
    "confidential": "AUTOCOMMERCE · CONFIDENTIAL",
    "pages": [
        (
            "Introduction and mandate",
            [
                (
                    "Objective",
                    "Assess the current AutoCommerce screens, convert the supplied annotations into "
                    "actionable requirements, and define a functional architecture for a distinctive, "
                    "simple and credible Figma/Make MVP.",
                ),
                (
                    "Positioning",
                    "AutoCommerce remains the product owner. The platform serves buyers, private or "
                    "professional sellers, and the AutoCommerce administration team. Cars.ca is the "
                    "functional coverage reference, without reproducing its identity or interface.",
                ),
                (
                    "Guiding principle",
                    "Simple at first sight, powerful on demand. Vehicles and the primary action remain "
                    "dominant; filters, analytics and expert capabilities progressively appear.",
                ),
            ],
        ),
        (
            "Observed scope and method",
            [
                (
                    "Evidence set",
                    "The evidence includes 18 annotated pages: home, Buy navigation, inventory, list "
                    "and gallery modes, vehicle detail, Buy Now, auctions, selling, authentication, "
                    "dashboard and the Add Vehicle assistant.",
                ),
                (
                    "Evaluation method",
                    "Each screen is assessed against user intent, visual hierarchy, cognitive load, "
                    "consistency, error exposure, accessibility, trust and responsive behaviour. "
                    "Findings are compared with Nielsen Norman Group heuristics and WCAG 2.2 AA.",
                ),
                (
                    "Limitations",
                    "Static captures do not reveal every behaviour, business rule or error state. "
                    "Recommendations therefore define requirements to validate in the prototype, "
                    "not a production-code audit.",
                ),
            ],
        ),
        (
            "Home and global navigation",
            [
                (
                    "Finding",
                    "The home experience relies on a strong black surface and automotive animation. "
                    "The identity is recognizable, but the value proposition, Buy/Sell paths and "
                    "next best action are not explicit enough.",
                ),
                (
                    "Recommendation",
                    "Retain the logo and established black, red and green while making white and light "
                    "grey the primary surfaces. Add central search, clear Buy/Sell entrances, trust "
                    "evidence and direct dashboard access.",
                ),
                (
                    "Target navigation",
                    "Compact persistent header: logo, Buy, Sell, Auctions, Services, Help, language "
                    "and account. On mobile, keep Search, Favourites, Sell and Account in a "
                    "thumb-accessible bottom navigation.",
                ),
            ],
        ),
        (
            "Inventory, search and filters",
            [
                (
                    "Primary issue",
                    "The annotation confirms that the static filter sidebar consumes too much space. "
                    "It reduces inventory visibility and competes with vehicle content. The Inventory "
                    "mega-menu also overwhelms the screen with simultaneous options.",
                ),
                (
                    "Functional solution",
                    "Replace the permanent sidebar with a collapsible desktop drawer and mobile bottom "
                    "sheet. Active criteria become removable chips. Closing filters preserves choices "
                    "and returns the full result width.",
                ),
                (
                    "Priority filters",
                    "Location, distance, make, model, year, price, mileage, body, fuel, transmission, "
                    "drivetrain, condition, seller and availability. Expert criteria remain under "
                    "“More filters.”",
                ),
            ],
        ),
        (
            "Results and vehicle detail",
            [
                (
                    "Exploration modes",
                    "List and gallery modes already exist. They must share the same data, filters and "
                    "states. A map can follow later. Changing view must never reset the search.",
                ),
                (
                    "Vehicle card",
                    "Primary image, year-make-model, trim, mileage, location, total price, confidence "
                    "badge, seller and Favourite/Compare actions. Buy Now, Make an Offer and Contact "
                    "must follow one consistent hierarchy.",
                ),
                (
                    "Detail page",
                    "Full gallery, price summary, history, condition, equipment, seller, indicative "
                    "financing, availability and actions. Critical information remains available "
                    "during scroll without masking the content.",
                ),
            ],
        ),
        (
            "Buy Now and auctions",
            [
                (
                    "Immediate purchase",
                    "Buy Now must show the attainable total price, mandatory fees, conditions, seller "
                    "identity and next step before commitment. Unattainable teaser pricing must not "
                    "be used.",
                ),
                (
                    "Offers",
                    "Making an offer requires an amount, expiry, summary and confirmation. Sent, viewed, "
                    "accepted, declined, expired and counter-offer states belong in the dashboard.",
                ),
                (
                    "Auctions",
                    "The current empty state should explain the service, show a calendar, enable alerts "
                    "and surface upcoming vehicles. A live auction needs server time, current bid, "
                    "increment, deposit, history, anti-sniping extension and confirmation.",
                ),
            ],
        ),
        (
            "Selling and listing lifecycle",
            [
                (
                    "Target lifecycle",
                    "Draft → quality check → pending approval → active → reserved → sold, expired or "
                    "declined. Every status must explain the reason, available action and next step.",
                ),
                (
                    "Dealer dashboard",
                    "Executive view of stock, listings to fix, hot leads, appointments, offers, response "
                    "times and sales. Inventory, Leads, Messages, Analytics, Team and Settings support "
                    "operational work.",
                ),
                (
                    "Productivity",
                    "Global search, saved filters, bulk actions, controlled duplication, CSV/API import, "
                    "lead assignment and activity history reduce dealer workload.",
                ),
            ],
        ),
        (
            "Add Vehicle assistant",
            [
                (
                    "Current structure",
                    "The captures show four steps: vehicle details, tyres, photos and damage, and "
                    "description. Many useful choices exist, but dense tags and long lists increase "
                    "effort and error risk.",
                ),
                (
                    "Recommended journey",
                    "1) VIN or manual entry; 2) verified specifications; 3) condition and tyres; "
                    "4) guided photos; 5) price and description; 6) preview. Autosave and progress "
                    "feedback support every step.",
                ),
                (
                    "Intelligent assistance",
                    "VIN pre-fills available data without blocking correction. Photo analysis detects "
                    "blur, duplicates and missing angles. Assisted copy summarizes verified facts "
                    "and remains fully editable.",
                ),
            ],
        ),
        (
            "Roles, dashboards and permissions",
            [
                (
                    "Buyer",
                    "Favourites, comparison, saved searches and alerts, messages, appointments, offers, "
                    "documents and purchase progress.",
                ),
                (
                    "Seller or dealer",
                    "Inventory, listing assistant, quality score, price recommendations, leads, tasks, "
                    "performance, team and exports.",
                ),
                (
                    "AutoCommerce",
                    "Accounts, dealers, moderation, publication rules, subscriptions, fraud, disputes, "
                    "quality, revenue, conversion, audit log and access control.",
                ),
                (
                    "Permissions",
                    "Owner administrator, manager, salesperson, marketing and read-only. Apply least "
                    "privilege, stronger authentication for sensitive actions and traceable history.",
                ),
            ],
        ),
        (
            "Target functional architecture",
            [
                (
                    "Experience layer",
                    "Responsive web and Figma prototype covering home, search, results, detail, account, "
                    "selling and administration. Components share tokens, states and accessibility rules.",
                ),
                (
                    "Business services",
                    "Catalogue, search, identity, organizations, listings, media, offers, auctions, leads, "
                    "messaging, notifications, pricing, moderation and analytics.",
                ),
                (
                    "MVP automation",
                    "Make simulates VIN intake, photo checks, draft generation, approval alerts and lead "
                    "notifications. Scenarios use fictional data and idempotent webhooks.",
                ),
                (
                    "Data and security",
                    "Separate users, organizations, vehicles, listings, media and events. Consent, "
                    "minimization, encryption, retention, access and deletion must be defined before "
                    "production.",
                ),
            ],
        ),
        (
            "Priorities, acceptance criteria and conclusion",
            [
                (
                    "MVP priority",
                    "Search and collapsible filters; list/gallery; vehicle detail; authentication; dealer "
                    "dashboard; VIN/photo/description assistant; AutoCommerce moderation; responsive and "
                    "accessible prototype.",
                ),
                (
                    "Key criteria",
                    "Filters close without losing choices; all actions work by keyboard; focus remains "
                    "visible; touch targets meet WCAG 2.2; errors are explicit; entry is saved; mandatory "
                    "prices are transparent; roles genuinely restrict access.",
                ),
                (
                    "Decision",
                    "AutoCommerce should not become another copy of an automotive marketplace. Its "
                    "difference is assisted selling: clearer inventory for buyers, faster publishing "
                    "for dealers and centralized quality control for AutoCommerce.",
                ),
            ],
        ),
    ],
    "sources": [
        "Supplied AutoCommerce capture — 18 annotated pages, 2026.",
        "AutoCommerce — public routes observed July 26, 2026: https://autocommerce.ca/.",
        "Nielsen Norman Group — 10 Usability Heuristics: https://www.nngroup.com/articles/ten-usability-heuristics/.",
        "Baymard Institute — Product Lists & Filtering UX: https://baymard.com/research/ecommerce-product-lists.",
        "Baymard Institute — Applied Filters Overview: https://baymard.com/blog/how-to-design-applied-filters.",
        "W3C Web Accessibility Initiative — WCAG 2.2: https://www.w3.org/TR/WCAG22/.",
        "Office of the Privacy Commissioner of Canada — Meaningful consent: https://www.priv.gc.ca/en/privacy-topics/privacy-laws-in-canada/the-personal-information-protection-and-electronic-documents-act-pipeda/p_principle/principles/p_consent/.",
        "Competition Bureau Canada — Drip pricing: https://competition-bureau.canada.ca/en/deceptive-marketing-practices/drip-pricing.",
        "Office of Consumer Affairs — Buying or leasing a vehicle: https://ised-isde.canada.ca/site/office-consumer-affairs/en/buying-and-leasing-big-ticket-items/buying-or-leasing-vehicle.",
    ],
}

FR["pages"].extend(
    [
        (
            "Audit du site vivant — constats vérifiés",
            [
                (
                    "Périmètre vérifié",
                    "Le 26 juillet 2026, les routes publiques /dashboard, /buy, /sell, /sold, "
                    "/contact, /register et /login ont été observées directement. Cette lecture "
                    "complète les captures sans prétendre couvrir les états authentifiés, les "
                    "réponses serveur, le clavier, le lecteur d’écran ou la performance.",
                ),
                (
                    "Tableau de bord public",
                    "Le tableau de bord expose Vehicle Listing, Buy Now et dix états transactionnels, "
                    "mais plusieurs valeurs sont remplacées par des points et la majorité des modules "
                    "portent la mention « coming soon ». L’interface promet un système opérationnel "
                    "plus large que ce qui est actuellement disponible.",
                ),
                (
                    "Parcours publics",
                    "Buy contient dix-sept familles de filtres et une seconde zone Make/Model. Sell "
                    "renvoie directement vers l’authentification. Sold ne fournit aucun contenu "
                    "interprétable. Contact duplique une même adresse courriel pour cinq fonctions. "
                    "Register demande un rôle sans en expliquer les conséquences.",
                ),
            ],
        ),
        (
            "Évaluation heuristique et charge cognitive",
            [
                (
                    "Visibilité de l’état",
                    "Les états vides, chargements et résultats n’expliquent pas toujours ce qui se "
                    "passe ni l’action suivante. Les libellés « ..... », les compteurs à zéro et "
                    "« coming soon » réduisent la confiance et ne différencient pas absence de donnée, "
                    "fonction inactive ou erreur.",
                ),
                (
                    "Correspondance avec le monde réel",
                    "Des termes comme D.Train, in if bid, pending delivery ou parked nécessitent un "
                    "vocabulaire normalisé et des explications contextuelles. Les acheteurs pensent "
                    "en budget, usage, état, localisation et confiance; les commerçants pensent en "
                    "stock, publication, prospects, négociation et rotation.",
                ),
                (
                    "Prévention des erreurs",
                    "Le changement de vue, la fermeture du filtre, l’enregistrement d’un brouillon et "
                    "la reprise après interruption doivent préserver l’état. Toute action financière, "
                    "publication, suppression ou changement de rôle exige un récapitulatif, une "
                    "confirmation proportionnée et un historique.",
                ),
            ],
        ),
        (
            "Audit d’accessibilité WCAG 2.2 AA",
            [
                (
                    "Risques observables",
                    "Les captures suggèrent de petites cibles, des contrastes subtils, une densité "
                    "élevée, des contrôles serrés et des informations parfois portées par la couleur. "
                    "Ces signaux ne constituent pas une certification; ils définissent les contrôles "
                    "techniques obligatoires.",
                ),
                (
                    "Contrôles requis",
                    "Tester le reflow à 320 CSS px, le zoom 400 %, la navigation clavier, l’ordre de "
                    "focus, le focus non masqué, les noms accessibles, les messages d’état, les erreurs, "
                    "la taille cible minimale, le contraste textuel et non textuel, ainsi que "
                    "l’authentification compatible avec gestionnaires de mots de passe et copier-coller.",
                ),
                (
                    "Seuil de livraison",
                    "WCAG 2.2 AA devient une porte de livraison. Les tests automatisés détectent une "
                    "partie des défauts; ils sont complétés par VoiceOver/NVDA, clavier seul, zoom, "
                    "mobile réel, contenu bilingue et personnes utilisatrices en situation de handicap.",
                ),
            ],
        ),
        (
            "Confiance, prix, confidentialité et sécurité",
            [
                (
                    "Transparence du prix",
                    "Le prix annoncé doit être atteignable avant taxes imposées par la loi. Les frais "
                    "obligatoires, transport, inspection, dépôt et conditions d’enchère doivent être "
                    "présentés avant l’engagement afin d’éviter le drip pricing et de faciliter la comparaison.",
                ),
                (
                    "VIN et données",
                    "Le VIN aide à vérifier l’historique et à préremplir les caractéristiques, mais il "
                    "doit rester corrigeable et traçable. Les photos, coordonnées, documents et données "
                    "d’usage nécessitent finalités explicites, consentement significatif, minimisation, "
                    "durées de conservation, accès restreint et suppression maîtrisée.",
                ),
                (
                    "Sécurité produit",
                    "Séparer comptes personnels et organisations, appliquer le moindre privilège, "
                    "protéger les médias, journaliser les actions sensibles, détecter les doublons et "
                    "soumettre les annonces à une modération proportionnée au risque.",
                ),
            ],
        ),
        (
            "Parcours cible — acheter un véhicule",
            [
                (
                    "Découvrir",
                    "Commencer par une recherche simple : marque/modèle, budget ou type d’usage et "
                    "localisation. Les critères avancés s’ouvrent dans un tiroir rétractable. Les "
                    "filtres actifs restent visibles sous forme de pastilles avec compteur de résultats.",
                ),
                (
                    "Évaluer",
                    "Comparer jusqu’à quatre véhicules sur prix total, kilométrage, état, historique, "
                    "équipements, localisation et vendeur. Les cartes montrent au moins trois images "
                    "ou un indicateur de galerie, les attributs discriminants et les signaux de confiance.",
                ),
                (
                    "Agir et suivre",
                    "Favori, alerte, contact, inspection, offre, Buy Now ou enchère sont contextualisés. "
                    "Le tableau de bord conserve les recherches, messages, rendez-vous, offres, documents "
                    "et étapes de transaction dans une timeline compréhensible.",
                ),
            ],
        ),
        (
            "Parcours cible — vendre et gérer un inventaire",
            [
                (
                    "Créer",
                    "Le commerçant scanne ou saisit le VIN, confirme les données décodées, documente "
                    "l’état, suit un guide photo, fixe le prix et vérifie une prévisualisation avant "
                    "soumission. Le brouillon est sauvegardé automatiquement.",
                ),
                (
                    "Publier",
                    "Un score de complétude explique les éléments manquants sans imposer une note opaque. "
                    "La génération assistée produit un brouillon factuel à partir des données confirmées; "
                    "le commerçant garde le contrôle éditorial et approuve la version publiée.",
                ),
                (
                    "Piloter",
                    "Le tableau de bord priorise les annonces à corriger, prospects sans réponse, offres "
                    "à traiter, rendez-vous et véhicules vieillissants. Des rôles, vues sauvegardées, "
                    "actions groupées et rapports soutiennent les équipes de plusieurs utilisateurs.",
                ),
            ],
        ),
        (
            "Architecture de l’information cible",
            [
                (
                    "Navigation publique",
                    "Acheter, Vendre, Enchères, Services, Aide, FR/EN et Compte. Le dashboard public "
                    "actuel devient une page d’accueil orientée vers les tâches et la confiance plutôt "
                    "qu’un inventaire de modules futurs.",
                ),
                (
                    "Espace acheteur",
                    "Accueil, Recherches, Favoris, Comparateur, Messages, Rendez-vous, Offres/Enchères, "
                    "Transactions, Documents et Paramètres.",
                ),
                (
                    "Espace commerçant",
                    "Vue d’ensemble, Inventaire, Ajouter un véhicule, Prospects, Messages, Transactions, "
                    "Analytique, Équipe et Paramètres. L’administration AutoCommerce ajoute Modération, "
                    "Organisations, Utilisateurs, Litiges, Qualité, Revenus et Journal d’audit.",
                ),
            ],
        ),
        (
            "Roadmap, mesure et gouvernance",
            [
                (
                    "Phase 0 — fondations",
                    "Valider objectifs, vocabulaire, responsabilités, analytics, risques, règles de prix, "
                    "données et protocole de recherche. Corriger les états publics incomplets et les "
                    "libellés les plus critiques.",
                ),
                (
                    "Phase 1 — prototype MVP",
                    "Prototyper recherche, filtres rétractables, résultats, fiche, compte, tableau de bord "
                    "commerçant et assistant VIN/photos/description. Tester mobile, desktop, clavier et "
                    "bilinguisme avant automatisation.",
                ),
                (
                    "Phase 2 — pilote",
                    "Connecter les scénarios Make à des données de démonstration, instrumenter le funnel "
                    "et piloter avec un petit groupe de commerçants. Mesurer réussite des tâches, temps "
                    "de publication, qualité des annonces, réponse aux prospects et conversion.",
                ),
            ],
        ),
        (
            "Registre des écrans et traçabilité",
            [
                (
                    "Principe",
                    "Chaque capture fournie possède un identifiant, un écran ou état, une santé initiale "
                    "et un constat principal. Ce registre évite de transformer une impression générale "
                    "en conclusion non traçable.",
                ),
                (
                    "Usage",
                    "Les futures maquettes Figma doivent référencer ces identifiants et le problème "
                    "résolu. Les écrans vivants sans équivalent dans le PDF sont ajoutés au registre "
                    "avec la date d’observation.",
                ),
            ],
        ),
        (
            "Conclusion stratégique",
            [
                (
                    "Décision recommandée",
                    "Faire d’AutoCommerce une plateforme de vente assistée, et non un catalogue de plus. "
                    "La différenciation vient d’une publication accélérée et fiable, d’une exploration "
                    "sans friction et d’un cockpit commerçant réellement opérationnel.",
                ),
                (
                    "Principe de conception",
                    "Clair d’abord, puissant à la demande : contenu véhicule dominant, filtres rétractables, "
                    "actions hiérarchisées, états explicites, automatisation contrôlée et preuves de confiance.",
                ),
                (
                    "Prochaine validation",
                    "Confirmer les règles métier, profils d’utilisateurs, états authentifiés, données disponibles "
                    "et priorités commerciales, puis matérialiser trois directions visuelles dans Figma avant de "
                    "figer le Design System et les scénarios Make.",
                ),
            ],
        ),
    ]
)

EN["pages"].extend(
    [
        (
            "Live-site audit — verified findings",
            [
                ("Verified scope", "On July 26, 2026, the public /dashboard, /buy, /sell, /sold, /contact, /register and /login routes were directly observed. This complements the supplied captures without claiming coverage of authenticated states, server responses, keyboard use, screen readers or performance."),
                ("Public dashboard", "The dashboard exposes Vehicle Listing, Buy Now and ten transaction states, yet several values are rendered as dots and most modules are marked “coming soon.” The interface promises a broader operational system than is currently available."),
                ("Public journeys", "Buy exposes seventeen filter families plus a second Make/Model area. Sell immediately redirects to authentication. Sold has no interpretable content. Contact repeats one email address for five functions. Register asks users to choose a role without explaining its implications."),
            ],
        ),
        (
            "Heuristic evaluation and cognitive load",
            [
                ("Visibility of system status", "Empty, loading and result states do not always explain what is happening or what users should do next. Dot placeholders, zero counters and “coming soon” do not distinguish no data, unavailable function or error."),
                ("Match with the real world", "Terms such as D.Train, in if bid, pending delivery and parked need normalized vocabulary and contextual explanations. Buyers think in budget, use, condition, location and trust; dealers think in stock, publishing, leads, negotiation and turn rate."),
                ("Error prevention", "View changes, closing filters, saving drafts and resuming after interruption must preserve state. Financial actions, publishing, deletion and role changes require a proportional review, confirmation and traceable history."),
            ],
        ),
        (
            "WCAG 2.2 AA accessibility audit",
            [
                ("Observable risks", "The captures suggest small targets, subtle contrast, high density, tight controls and information sometimes carried by colour. These signals are not a certification; they define mandatory technical tests."),
                ("Required checks", "Test reflow at 320 CSS px, 400% zoom, keyboard order, visible and unobscured focus, accessible names, status messages, errors, minimum target size, text and non-text contrast, and authentication compatible with password managers and copy/paste."),
                ("Release threshold", "WCAG 2.2 AA becomes a release gate. Automated checks cover only part of the risk and must be complemented by VoiceOver/NVDA, keyboard-only use, zoom, real mobile devices, bilingual content and participants with disabilities."),
            ],
        ),
        (
            "Trust, pricing, privacy and security",
            [
                ("Price transparency", "The represented price must be attainable before government-imposed taxes. Mandatory fees, transportation, inspection, deposits and auction conditions must appear before commitment to prevent drip pricing and support comparison."),
                ("VIN and personal data", "VIN data can verify history and prefill specifications, but must remain correctable and traceable. Photos, contact details, documents and usage data require explicit purposes, meaningful consent, minimization, retention limits, restricted access and controlled deletion."),
                ("Product security", "Separate people from organizations, enforce least privilege, protect media, log sensitive actions, detect duplicates and apply risk-based listing moderation."),
            ],
        ),
        (
            "Target journey — buying a vehicle",
            [
                ("Discover", "Start with make/model, budget or intended use and location. Advanced criteria open in a collapsible drawer. Applied filters remain visible as removable chips with result count."),
                ("Evaluate", "Compare up to four vehicles across total price, mileage, condition, history, equipment, location and seller. Cards show at least three images or a gallery indicator, differentiating attributes and trust signals."),
                ("Act and track", "Favourite, alert, contact, inspection, offer, Buy Now and auction actions are contextual. The dashboard retains searches, messages, appointments, offers, documents and transaction steps in a clear timeline."),
            ],
        ),
        (
            "Target journey — selling and inventory operations",
            [
                ("Create", "The dealer scans or enters the VIN, confirms decoded data, records condition, follows guided photography, sets price and reviews a preview before submission. Drafts autosave."),
                ("Publish", "A completeness score explains missing elements without imposing an opaque grade. Assisted generation creates a factual draft from confirmed data; the dealer retains editorial control and approves the published version."),
                ("Operate", "The dashboard prioritizes listings to fix, unanswered leads, offers, appointments and ageing inventory. Roles, saved views, bulk actions and reporting support multi-user teams."),
            ],
        ),
        (
            "Target information architecture",
            [
                ("Public navigation", "Buy, Sell, Auctions, Services, Help, FR/EN and Account. The current public dashboard becomes a task- and trust-oriented homepage rather than an inventory of future modules."),
                ("Buyer workspace", "Home, Searches, Favourites, Compare, Messages, Appointments, Offers/Auctions, Transactions, Documents and Settings."),
                ("Dealer workspace", "Overview, Inventory, Add Vehicle, Leads, Messages, Transactions, Analytics, Team and Settings. AutoCommerce administration adds Moderation, Organizations, Users, Disputes, Quality, Revenue and Audit Log."),
            ],
        ),
        (
            "Roadmap, measurement and governance",
            [
                ("Phase 0 — foundations", "Validate objectives, vocabulary, responsibilities, analytics, risks, pricing rules, data and research protocol. Correct incomplete public states and critical labels."),
                ("Phase 1 — MVP prototype", "Prototype search, collapsible filters, results, vehicle detail, account, dealer dashboard and VIN/photo/description assistant. Test mobile, desktop, keyboard and bilingual content before automation."),
                ("Phase 2 — pilot", "Connect Make scenarios to demonstration data, instrument the funnel and pilot with a small dealer cohort. Measure task success, publishing time, listing quality, lead response and conversion."),
            ],
        ),
        (
            "Screen register and traceability",
            [
                ("Principle", "Every supplied capture receives an identifier, screen or state, initial health rating and primary finding. This register prevents broad impressions from becoming untraceable conclusions."),
                ("Use", "Future Figma frames should reference these identifiers and the issue solved. Live screens without a PDF equivalent are added with their observation date."),
            ],
        ),
        (
            "Strategic conclusion",
            [
                ("Recommended decision", "Position AutoCommerce as an assisted-selling platform, not another catalogue. Differentiation comes from fast and reliable publishing, frictionless exploration and a genuinely operational dealer cockpit."),
                ("Design principle", "Clear first, powerful on demand: vehicle content dominates, filters collapse, actions are hierarchical, states are explicit, automation remains controlled and trust is evidenced."),
                ("Next validation", "Confirm business rules, user profiles, authenticated states, available data and commercial priorities, then materialize three visual directions in Figma before locking the Design System and Make scenarios."),
            ],
        ),
    ]
)

MATRIX_DATA = {
    "FR": {
        "Évaluation heuristique et charge cognitive": (
            ["ID", "Sévérité", "Problème", "Action prioritaire"],
            [
                ["P01", "Critique", "Filtres dominants et redondants", "Tiroir rétractable + pastilles actives"],
                ["P02", "Critique", "États vides ou ambigus", "État, cause et prochaine action"],
                ["P03", "Élevée", "Sell bloqué sans valeur expliquée", "Prévisualiser le parcours avant connexion"],
                ["P04", "Élevée", "Rôles d’inscription non expliqués", "Définir profils, droits et conséquences"],
                ["P05", "Élevée", "Terminologie incohérente", "Lexique produit bilingue"],
                ["P06", "Élevée", "Dashboard orienté modules", "Cockpit orienté tâches et exceptions"],
            ],
        ),
        "Audit d’accessibilité WCAG 2.2 AA": (
            ["Critère", "Risque à contrôler", "Validation"],
            [
                ["1.4.10", "Reflow des filtres et tableaux", "320 CSS px / zoom 400 %"],
                ["1.4.3 / 1.4.11", "Contrastes texte et composants", "Mesure sur tokens réels"],
                ["2.4.7 / 2.4.11", "Focus visible ou masqué", "Parcours clavier complet"],
                ["2.5.8", "Petites cibles tactiles", "24 × 24 CSS px minimum; viser 44"],
                ["3.3.1 / 3.3.3", "Erreurs insuffisamment reliées", "Résumé + message adjacent"],
                ["3.3.8", "Charge cognitive à la connexion", "Gestionnaire de mots de passe, copier-coller"],
            ],
        ),
        "Roadmap, mesure et gouvernance": (
            ["Horizon", "Livrable", "Gate de décision"],
            [
                ["0–4 sem.", "Diagnostic vivant, données, vocabulaire", "Périmètre et risques validés"],
                ["5–12 sem.", "Prototype Figma responsive FR/EN", "Tests tâches et WCAG passés"],
                ["13–20 sem.", "Scénarios Make + pilote commerçants", "Valeur et sécurité démontrées"],
                ["6–12 mois", "Plateforme progressive", "KPI, conformité et rentabilité"],
            ],
        ),
        "Registre des écrans et traçabilité": (
            ["Pages", "Surface", "Santé", "Constat principal"],
            [
                ["1", "Accueil", "Fragile", "Valeur et actions peu hiérarchisées"],
                ["2–6", "Buy / inventaire / galerie", "Critique", "Filtres dominants; résultats concurrencés"],
                ["7–12", "Sell / états", "Critique", "États vides et progression peu explicites"],
                ["13", "Après connexion", "Fragile", "Contexte et prochaine action faibles"],
                ["14–18", "Add Vehicle", "Fragile", "Densité, listes longues et progression"],
            ],
        ),
    },
    "EN": {
        "Heuristic evaluation and cognitive load": (
            ["ID", "Severity", "Problem", "Priority action"],
            [
                ["P01", "Critical", "Dominant and duplicated filters", "Collapsible drawer + active chips"],
                ["P02", "Critical", "Empty or ambiguous states", "State, cause and next action"],
                ["P03", "High", "Sell blocked without value explanation", "Preview journey before login"],
                ["P04", "High", "Registration roles unexplained", "Define profiles, rights and consequences"],
                ["P05", "High", "Inconsistent terminology", "Bilingual product lexicon"],
                ["P06", "High", "Module-oriented dashboard", "Task- and exception-oriented cockpit"],
            ],
        ),
        "WCAG 2.2 AA accessibility audit": (
            ["Criterion", "Risk to test", "Validation"],
            [
                ["1.4.10", "Filter and table reflow", "320 CSS px / 400% zoom"],
                ["1.4.3 / 1.4.11", "Text and component contrast", "Measure actual design tokens"],
                ["2.4.7 / 2.4.11", "Visible or obscured focus", "Complete keyboard journey"],
                ["2.5.8", "Small touch targets", "24 × 24 CSS px minimum; aim for 44"],
                ["3.3.1 / 3.3.3", "Errors not connected to fields", "Summary + adjacent message"],
                ["3.3.8", "Cognitive burden at login", "Password manager and copy/paste"],
            ],
        ),
        "Roadmap, measurement and governance": (
            ["Horizon", "Deliverable", "Decision gate"],
            [
                ["0–4 weeks", "Live diagnostic, data, vocabulary", "Scope and risks validated"],
                ["5–12 weeks", "Responsive bilingual Figma prototype", "Task and WCAG tests passed"],
                ["13–20 weeks", "Make scenarios + dealer pilot", "Value and security demonstrated"],
                ["6–12 months", "Progressive platform", "KPI, compliance and profitability"],
            ],
        ),
        "Screen register and traceability": (
            ["Pages", "Surface", "Health", "Primary finding"],
            [
                ["1", "Home", "Fragile", "Value and actions lack hierarchy"],
                ["2–6", "Buy / inventory / gallery", "Critical", "Filters dominate vehicle results"],
                ["7–12", "Sell / states", "Critical", "Empty states and progress unclear"],
                ["13", "Post-login", "Fragile", "Weak context and next action"],
                ["14–18", "Add Vehicle", "Fragile", "Density, long lists and progress"],
            ],
        ),
    },
}

UX_CRITERIA_FR = [
    ("Hiérarchie & clarté", "L’objectif, le contenu et l’action principale sont immédiatement compréhensibles."),
    ("Navigation & architecture", "La position, les choix et le chemin de retour sont prévisibles."),
    ("Efficacité & charge cognitive", "L’effort, la densité et le nombre de décisions sont proportionnés à la tâche."),
    ("Cohérence & standards", "Les libellés, composants et comportements restent constants."),
    ("Feedback & visibilité de l’état", "Le système explique chargement, résultat, progression, succès et échec."),
    ("Prévention & récupération", "Les erreurs sont évitées, expliquées et réparables sans perte."),
    ("Accessibilité", "Le contenu vise WCAG 2.2 AA, clavier, zoom, reflow et technologies d’assistance."),
    ("Responsive & tactile", "La priorité et les contrôles restent utilisables sur mobile."),
    ("Confiance & contenu", "Prix, vendeur, données, conditions et prochaines étapes sont transparents."),
]

UX_CRITERIA_EN = [
    ("Hierarchy & clarity", "Purpose, content and primary action are immediately understandable."),
    ("Navigation & architecture", "Location, choices and return path are predictable."),
    ("Efficiency & cognitive load", "Effort, density and decisions are proportional to the task."),
    ("Consistency & standards", "Labels, components and behaviours remain consistent."),
    ("Feedback & system status", "Loading, results, progress, success and failure are explained."),
    ("Prevention & recovery", "Errors are prevented, explained and recoverable without data loss."),
    ("Accessibility", "The experience targets WCAG 2.2 AA, keyboard, zoom, reflow and assistive technology."),
    ("Responsive & touch", "Priorities and controls remain usable on mobile."),
    ("Trust & content", "Price, seller, data, conditions and next steps are transparent."),
]

SCREEN_AUDIT_FR = [
    ("Écran 01 — Accueil", 1, "Découvrir la proposition AutoCommerce et choisir Acheter ou Vendre.", [
        ("Hiérarchie", "2/4", "Vidéo et noir dominants; promesse et CTA secondaires.", "Orientation lente", "Recherche + Acheter/Vendre au-dessus de la ligne de flottaison."),
        ("Navigation", "2/4", "Menu visible mais modèle produit non expliqué.", "Choix incertain", "Navigation publique courte avec états actifs et langue."),
        ("Contenu", "1/4", "Peu de preuves de confiance et de valeur métier.", "Faible crédibilité", "Prix transparent, inspection, VIN, commerçants vérifiés."),
        ("Accessibilité", "À tester", "Animation, contraste et focus non vérifiables sur capture.", "Barrières possibles", "Pause animation, contraste mesuré, clavier et reflow."),
        ("Mobile", "2/4", "Composition visuelle lourde pour un petit écran.", "CTA repoussés", "Hero léger et navigation basse orientée tâches."),
    ]),
    ("Écran 02 — Menu Buy / Inventory", 2, "Accéder rapidement au bon mode d’achat.", [
        ("Hiérarchie", "2/4", "Inventory, Buy Now et Live Auction ont un poids proche.", "Priorité ambiguë", "Définir Inventory comme entrée principale; autres modes secondaires."),
        ("Architecture", "2/4", "Menu superposé large et catégories peu expliquées.", "Charge de décision", "Méga-menu compact avec descriptions et raccourcis."),
        ("Cohérence", "2/4", "Terminologie mixte et capitalisation irrégulière.", "Compréhension réduite", "Lexique FR/EN et règles de nommage."),
        ("Feedback", "1/4", "Pas d’indication de disponibilité ni volumes.", "Attentes trompeuses", "Afficher compteurs réels et disponibilité."),
        ("Mobile", "1/4", "Menu de grande largeur difficile à transposer.", "Navigation fragile", "Panneau plein écran avec 3 choix maximum."),
    ]),
    ("Écran 03 — Inventaire et fiche véhicule", 3, "Parcourir les véhicules et comprendre un véhicule sélectionné.", [
        ("Hiérarchie", "2/4", "Liste, badges, filtres et fiche se concurrencent.", "Balayage lent", "Grille maître-détail avec prix et véhicule dominants."),
        ("Efficacité", "2/4", "Information dense et répétitive dans les cartes.", "Comparaison difficile", "Attributs discriminants constants et comparaison dédiée."),
        ("Feedback", "2/4", "Sélection visible mais relation liste/fiche faible.", "Perte de contexte", "État sélectionné fort, compteur et URL persistante."),
        ("Confiance", "2/4", "Prix, état, vendeur et inspection manquent de structure.", "Décision risquée", "Bloc confiance normalisé et prix total atteignable."),
        ("Responsive", "1/4", "Trois zones simultanées ne peuvent tenir sur mobile.", "Reflow critique", "Liste puis fiche; retour conservant position et filtres."),
    ]),
    ("Écran 04 — Filtres latéraux", 4, "Réduire l’inventaire sans perdre la visibilité des véhicules.", [
        ("Hiérarchie", "1/4", "Le filtre permanent occupe une part importante de l’écran.", "Contenu masqué", "Tiroir rétractable; inventaire pleine largeur à la fermeture."),
        ("Efficacité", "2/4", "Nombreux critères visibles en même temps.", "Charge cognitive", "Filtres prioritaires puis « Plus de critères »."),
        ("Feedback", "1/4", "Impact des choix et filtres actifs peu synthétisés.", "Désorientation", "Pastilles actives, compteur et suppression individuelle."),
        ("Prévention", "2/4", "Risque de remise à zéro ou perte au changement de vue.", "Travail perdu", "Persistance, URL partageable et Reset confirmé."),
        ("Mobile", "1/4", "Panneau latéral non adapté au pouce.", "Blocage", "Bottom sheet avec bouton « Voir X véhicules »."),
    ]),
    ("Écran 05 — Buy Now / mode liste", 5, "Évaluer les véhicules disponibles à l’achat immédiat.", [
        ("Hiérarchie", "2/4", "Badges, actions et métadonnées sont très compacts.", "CTA confus", "Un CTA primaire; prix et disponibilité avant détails."),
        ("Comparaison", "2/4", "Attributs variables entre lignes.", "Évaluation inégale", "Gabarit stable : prix, km, état, lieu, vendeur."),
        ("Confiance", "1/4", "Prix total et frais obligatoires non démontrés.", "Risque légal", "Ventilation avant engagement et conditions accessibles."),
        ("Accessibilité", "À tester", "Petits textes et contrôles suggérés.", "Lecture/action difficiles", "Base 16 px, cibles et contrastes mesurés."),
        ("Responsive", "2/4", "Liste dense et actions horizontales.", "Défilement excessif", "Carte verticale compacte et CTA fixe non masquant."),
    ]),
    ("Écran 06 — Galerie et Live Auctions", 6, "Explorer visuellement et comprendre les enchères.", [
        ("Hiérarchie", "2/4", "Galerie utile, mais plusieurs badges concurrencent les images.", "Lecture morcelée", "Image, prix, modèle puis badges secondaires."),
        ("Cohérence", "2/4", "Liste et galerie ne démontrent pas une continuité d’état.", "Filtres perdus", "Même source de données, tri et sélection persistants."),
        ("État vide", "1/4", "Live Auctions affiche peu ou pas d’explication.", "Impression de panne", "Calendrier, règles, alertes et véhicules à venir."),
        ("Confiance", "1/4", "Dépôt, incrément, heure et prolongation non visibles.", "Engagement incertain", "Règles d’enchère et heure serveur persistantes."),
        ("Mobile", "2/4", "La galerie s’adapte mieux, mais enchères exigent temps réel.", "Erreur de mise", "Confirmation forte et latence explicitée."),
    ]),
    ("Écran 07 — Sell / barrière de connexion", 7, "Comprendre comment vendre avant de créer un compte.", [
        ("Valeur", "1/4", "Message « login first » sans bénéfices ni étapes.", "Abandon", "Expliquer VIN, photos, délai, prix et contrôle qualité."),
        ("Hiérarchie", "2/4", "Signup et Login sont présents mais sans recommandation.", "Choix inutile", "CTA « Commencer une annonce » puis connexion contextuelle."),
        ("Aide", "1/4", "Aucun aperçu des prérequis.", "Surprise après compte", "Checklist et temps estimé avant authentification."),
        ("Accessibilité", "À tester", "État de focus et alternatives non observables.", "Barrières possibles", "Authentification WCAG 3.3.8 et messages explicites."),
        ("Confiance", "1/4", "Utilisation des photos/VIN non expliquée.", "Réticence aux données", "Notice de confidentialité au point de collecte."),
    ]),
    ("Écran 08 — Entrée vendeur / état initial", 8, "Démarrer ou reprendre une mise en vente.", [
        ("Clarté", "2/4", "Écran très vide et contexte limité.", "Prochaine action faible", "Titre, bénéfice, progression et reprise de brouillon."),
        ("Feedback", "1/4", "Aucun statut de sauvegarde ou d’avancement.", "Incertitude", "Autosave horodaté et état du brouillon."),
        ("Prévention", "2/4", "Peu d’aide sur les données requises.", "Erreurs tardives", "Validation progressive et exemples."),
        ("Cohérence", "2/4", "Style déconnecté des surfaces Buy.", "Produit fragmenté", "Design System commun et patterns de formulaire."),
        ("Mobile", "2/4", "Peu de contenu, mais commandes à valider.", "Cibles incertaines", "Une intention par écran et boutons pleine largeur."),
    ]),
    ("Écran 09 — Annonces Pending", 9, "Comprendre quelles annonces attendent et pourquoi.", [
        ("État", "1/4", "Pending est un libellé sans explication.", "Support accru", "Cause, propriétaire, date et délai estimé."),
        ("Action", "1/4", "Aucune prochaine action clairement visible.", "Blocage", "Corriger, compléter, contacter ou attendre."),
        ("Tableau", "2/4", "Colonnes et densité peu hiérarchisées.", "Balayage lent", "Colonnes essentielles, tri et actions en fin de ligne."),
        ("Confiance", "2/4", "Processus de validation non expliqué.", "Perte de contrôle", "Timeline de modération et journal."),
        ("Responsive", "1/4", "Table large et zones vides.", "Scroll bidirectionnel", "Cartes mobiles ou colonnes prioritaires."),
    ]),
    ("Écran 10 — Annonces Finished", 10, "Retrouver et analyser les annonces terminées.", [
        ("Clarté", "1/4", "Finished mélange potentiellement vendu, expiré et refusé.", "Historique ambigu", "Statuts distincts et définitions."),
        ("Efficacité", "2/4", "Peu de moyens visibles pour filtrer ou réutiliser.", "Travail répétitif", "Filtres, duplication contrôlée et export."),
        ("Feedback", "1/4", "Résultat final et motif non saillants.", "Apprentissage impossible", "Résultat, date, prix et raison."),
        ("Analytique", "1/4", "Aucun lien vers performance ou prospects.", "Valeur perdue", "Vues, contacts, délai, conversion et recommandations."),
        ("Responsive", "2/4", "État vide simple mais peu utile.", "Faible utilité mobile", "Résumé par carte et action contextuelle."),
    ]),
    ("Écran 11 — Liste vendeur / recherche", 11, "Retrouver et gérer une annonce précise.", [
        ("Recherche", "2/4", "Contrôles présents mais peu explicites.", "Temps de recherche", "Recherche globale VIN, stock, marque et statut."),
        ("Hiérarchie", "2/4", "Résultat et détails utilisent peu l’espace.", "Efficacité réduite", "Densité adaptative et colonnes configurables."),
        ("Actions", "1/4", "Actions principales peu visibles ou absentes.", "Gestion bloquée", "Éditer, publier, suspendre, dupliquer et archiver."),
        ("Prévention", "1/4", "Risque d’actions sensibles sans distinction.", "Erreur opérationnelle", "Séparer actions courantes et dangereuses."),
        ("Accessibilité", "À tester", "Table et menus nécessitent clavier/lecteur d’écran.", "Navigation complexe", "En-têtes, noms accessibles et focus géré."),
    ]),
    ("Écran 12 — États vendeur / variation", 12, "Passer d’un statut d’inventaire à un autre.", [
        ("Architecture", "2/4", "États répartis dans des onglets peu différenciés.", "Modèle mental lourd", "Cycle de vie explicite et filtres par statut."),
        ("Cohérence", "2/4", "Compteurs et libellés variables.", "Confiance réduite", "Taxonomie unique et règles de transition."),
        ("Feedback", "1/4", "Changements de statut non expliqués.", "Support et erreurs", "Toast utile + journal + état précédent."),
        ("Prévention", "2/4", "Transitions potentiellement irréversibles.", "Perte de vente", "Confirmation selon risque et possibilité d’annuler."),
        ("Mobile", "1/4", "Onglets et tables dépassent facilement.", "Reflow fragile", "Filtres de statut et cartes responsives."),
    ]),
    ("Écran 13 — Après connexion / menu compte", 13, "S’orienter immédiatement après authentification.", [
        ("Orientation", "2/4", "Menu ouvert mais rôle et contexte peu saillants.", "Mauvaise destination", "Nom, organisation, rôle et tâche suivante."),
        ("Architecture", "2/4", "Accès compte et produit semblent séparés.", "Navigation fragmentée", "Shell authentifié unique."),
        ("Sécurité", "2/4", "Actions et permissions non visibles.", "Erreur de rôle", "Rôles explicites et actions conditionnelles."),
        ("Feedback", "2/4", "Connexion réussie sans synthèse opérationnelle.", "Temps perdu", "Dashboard personnalisé par rôle."),
        ("Accessibilité", "À tester", "Menu et focus après ouverture non vérifiables.", "Piège clavier possible", "Focus initial, Escape et retour au déclencheur."),
    ]),
    ("Écran 14 — Add Vehicle / détails", 14, "Identifier correctement le véhicule.", [
        ("Clarté", "2/4", "Nombreux champs dans une grande modale.", "Effort élevé", "VIN d’abord, puis données préremplies et confirmables."),
        ("Progression", "2/4", "Étapes présentes mais état complet/incomplet faible.", "Omissions", "Stepper avec statuts et erreurs par étape."),
        ("Prévention", "2/4", "Relations Make/Model/Trim susceptibles d’erreur.", "Données incohérentes", "Dépendances, validation serveur et correction manuelle."),
        ("Feedback", "1/4", "Sauvegarde et validation peu visibles.", "Perte perçue", "Autosave, horodatage et messages adjacents."),
        ("Mobile", "1/4", "Modal dense et colonnes.", "Saisie difficile", "Page dédiée, une section courte à la fois."),
    ]),
    ("Écran 15 — Add Vehicle / pneus et couleurs", 15, "Décrire l’état et les caractéristiques visuelles.", [
        ("Hiérarchie", "2/4", "Palettes et choix nombreux dominent la tâche.", "Décision lente", "Choix récents/prioritaires puis recherche."),
        ("Accessibilité", "1/4", "Couleur possiblement utilisée seule.", "Exclusion", "Nom textuel, sélection, contraste et motif."),
        ("Cohérence", "2/4", "Pneus, couleurs et attributs utilisent des patterns variés.", "Apprentissage", "Radio, checkbox et sélecteur normalisés."),
        ("Prévention", "2/4", "Valeurs techniques sans aide.", "Saisie incorrecte", "Exemples, unités et « Inconnu »."),
        ("Mobile", "1/4", "Grande palette et nombreuses options.", "Défilement", "Recherche, accordéons et résumé persistant."),
    ]),
    ("Écran 16 — Add Vehicle / équipements", 16, "Sélectionner les options et équipements pertinents.", [
        ("Charge cognitive", "1/4", "Longue liste de pastilles sans regroupement clair.", "Fatigue et omissions", "Catégories, recherche et suggestions issues du VIN."),
        ("Hiérarchie", "2/4", "Toutes les options ont le même poids.", "Balayage lent", "Équipements clés en premier; avancés repliés."),
        ("Feedback", "2/4", "Sélection visible mais résumé global faible.", "Doute", "Compteur et résumé modifiable."),
        ("Accessibilité", "1/4", "Petites cibles et nombreux toggles possibles.", "Interaction difficile", "Cibles ≥24 px, viser 44, focus et état textuel."),
        ("Performance", "À tester", "Grand nombre de composants interactifs.", "Lenteur mobile", "Rendu progressif et virtualisation si nécessaire."),
    ]),
    ("Écran 17 — Add Vehicle / photos et dommages", 17, "Prouver l’état du véhicule avec des médias fiables.", [
        ("Guidage", "2/4", "Angles et dommages présents mais chemin peu pédagogique.", "Photos incomplètes", "Checklist illustrée avant téléversement."),
        ("Feedback", "1/4", "Qualité, progression et erreurs d’upload peu visibles.", "Reprise difficile", "Progression par fichier, reprise et statut."),
        ("Prévention", "2/4", "Flou, doublons et formats peuvent être découverts tard.", "Rejet tardif", "Contrôle local + analyse assistée explicable."),
        ("Confidentialité", "1/4", "Usage, conservation et métadonnées non expliqués.", "Risque de confiance", "Consentement, retrait EXIF et politique de rétention."),
        ("Mobile", "2/4", "Capture mobile naturelle mais réseau variable.", "Échec en mobilité", "Compression, reprise et mode faible débit."),
    ]),
    ("Écran 18 — Add Vehicle / description et soumission", 18, "Finaliser une annonce exacte avant publication.", [
        ("Clarté", "2/4", "Description et actions finales manquent de résumé global.", "Erreurs restantes", "Review page avec liens vers chaque correction."),
        ("Assistance IA", "À cadrer", "Génération rapide évoquée sans garde-fous visibles.", "Contenu faux", "Brouillon factuel, sources des champs et approbation humaine."),
        ("Prévention", "1/4", "Publication potentiellement irréversible.", "Annonce incorrecte", "Prévisualisation, attestation et confirmation."),
        ("Feedback", "1/4", "Statut après envoi et délai peu explicites.", "Incertitude", "Reçu, timeline de modération et notification."),
        ("Confiance", "2/4", "Responsabilité du contenu non explicitée.", "Litige", "Historique des versions et propriétaire de chaque modification."),
    ]),
]

SCREEN_AUDIT_EN = [
    (title.replace("Écran", "Screen").replace("Accueil", "Home").replace("Menu Buy / Inventory", "Buy / Inventory menu")
     .replace("Inventaire et fiche véhicule", "Inventory and vehicle detail").replace("Filtres latéraux", "Filter sidebar")
     .replace("Buy Now / mode liste", "Buy Now / list mode").replace("Galerie et Live Auctions", "Gallery and Live Auctions")
     .replace("Sell / barrière de connexion", "Sell / login barrier").replace("Entrée vendeur / état initial", "Seller entry / initial state")
     .replace("Annonces Pending", "Pending listings").replace("Annonces Finished", "Finished listings")
     .replace("Liste vendeur / recherche", "Seller list / search").replace("États vendeur / variation", "Seller status variation")
     .replace("Après connexion / menu compte", "Post-login / account menu").replace("Add Vehicle / détails", "Add Vehicle / details")
     .replace("Add Vehicle / pneus et couleurs", "Add Vehicle / tyres and colours").replace("Add Vehicle / équipements", "Add Vehicle / equipment")
     .replace("Add Vehicle / photos et dommages", "Add Vehicle / photos and damage").replace("Add Vehicle / description et soumission", "Add Vehicle / description and submission"),
     page,
     {
         1: "Discover the AutoCommerce proposition and choose Buy or Sell.",
         2: "Reach the appropriate buying mode quickly.",
         3: "Browse vehicles and understand the selected vehicle.",
         4: "Narrow inventory without losing vehicle visibility.",
         5: "Evaluate vehicles available for immediate purchase.",
         6: "Explore visually and understand auctions.",
         7: "Understand how selling works before creating an account.",
         8: "Start or resume a listing.",
         9: "Understand which listings are pending and why.",
         10: "Find and analyse completed listings.",
         11: "Find and manage a specific listing.",
         12: "Move between inventory statuses.",
         13: "Orient immediately after authentication.",
         14: "Identify the vehicle correctly.",
         15: "Describe condition and visual characteristics.",
         16: "Select relevant options and equipment.",
         17: "Evidence vehicle condition with reliable media.",
         18: "Finalize an accurate listing before publication.",
     }[page],
     [(criterion, score, evidence, impact, recommendation) for criterion, score, evidence, impact, recommendation in rows])
    for title, page, _, rows in SCREEN_AUDIT_FR
]

def install_detailed_audit(language: dict, screen_data: list, criteria: list) -> None:
    methodology_title = (
        "Cadre d’évaluation UX/UI et échelle de notation"
        if language["code"] == "FR"
        else "UX/UI evaluation framework and scoring scale"
    )
    methodology_sections = [
        (
            "Méthode" if language["code"] == "FR" else "Method",
            (
                "Chaque écran est évalué contre neuf dimensions constantes. La note 1 signifie "
                "critique, 2 fragile, 3 acceptable et 4 robuste. « À tester » distingue les "
                "propriétés impossibles à certifier depuis une capture."
                if language["code"] == "FR"
                else "Every screen is evaluated against nine consistent dimensions. Score 1 means "
                "critical, 2 fragile, 3 acceptable and 4 robust. “To test” separates properties "
                "that cannot be certified from a static capture."
            ),
        ),
        (
            "Règle de preuve" if language["code"] == "FR" else "Evidence rule",
            (
                "Une observation visible est formulée comme constat. Un comportement non visible "
                "devient un risque ou un contrôle requis. La sévérité combine fréquence, blocage, "
                "conséquence commerciale, accessibilité et possibilité de récupération."
                if language["code"] == "FR"
                else "Visible evidence is stated as a finding. Non-visible behaviour becomes a risk "
                "or required check. Severity combines frequency, blockage, business consequence, "
                "accessibility and recoverability."
            ),
        ),
    ]
    insert_at = 2
    language["pages"][insert_at:insert_at] = [(methodology_title, methodology_sections)]
    language["pages"][insert_at + 1:insert_at + 1] = [
        (
            title,
            [
                (
                    "Objectif utilisateur" if language["code"] == "FR" else "User objective",
                    objective,
                ),
                (
                    "Diagnostic" if language["code"] == "FR" else "Diagnosis",
                    (
                        "La matrice ci-dessous relie le constat visible à son impact et à une "
                        "recommandation vérifiable dans le prototype."
                        if language["code"] == "FR"
                        else "The matrix below connects visible evidence to user impact and a "
                        "recommendation that can be verified in the prototype."
                    ),
                ),
            ],
        )
        for title, _, objective, _ in screen_data
    ]
    criteria_headers = (
        ["Critère", "Définition"] if language["code"] == "FR" else ["Criterion", "Definition"]
    )
    MATRIX_DATA[language["code"]][methodology_title] = (
        criteria_headers,
        [[name, definition] for name, definition in criteria],
    )
    for title, page, _, rows in screen_data:
        MATRIX_DATA[language["code"]][title] = (
            (
                ["Critère", "Score", "Preuve", "Impact", "Recommandation"]
                if language["code"] == "FR"
                else ["Criterion", "Score", "Evidence", "Impact", "Recommendation"]
            ),
            [list(row) for row in rows],
        )

install_detailed_audit(FR, SCREEN_AUDIT_FR, UX_CRITERIA_FR)
install_detailed_audit(EN, SCREEN_AUDIT_EN, UX_CRITERIA_EN)

EN_AUDIT_POINTS = {
    1: [
        ["Hierarchy", "2/4", "Video and dark surfaces dominate the proposition.", "Slow orientation", "Place search and Buy/Sell choices above the fold."],
        ["Navigation", "2/4", "Menu is visible but the product model is unexplained.", "Uncertain choice", "Use a short public navigation with clear active states."],
        ["Content", "1/4", "Trust and marketplace value are weakly evidenced.", "Low credibility", "Show inspection, VIN, verified seller and price promises."],
        ["Accessibility", "To test", "Animation, contrast and focus cannot be certified.", "Potential barriers", "Pause control, measured contrast, keyboard and reflow tests."],
        ["Mobile", "2/4", "The visual composition is heavy for a small screen.", "Actions pushed down", "Use a lightweight hero and task-oriented bottom navigation."],
    ],
    2: [
        ["Hierarchy", "2/4", "Inventory, Buy Now and Live Auction have similar weight.", "Ambiguous priority", "Make Inventory primary and describe secondary modes."],
        ["Architecture", "2/4", "Large overlay menu with unexplained categories.", "Decision burden", "Use a compact menu with descriptions and shortcuts."],
        ["Consistency", "2/4", "Mixed terminology and capitalization.", "Reduced comprehension", "Adopt a bilingual product lexicon."],
        ["Feedback", "1/4", "Availability and volumes are not communicated.", "Misleading expectations", "Show real counts and availability."],
        ["Mobile", "1/4", "Wide menu is difficult to transpose.", "Fragile navigation", "Use a full-screen panel with three primary choices."],
    ],
    3: [
        ["Hierarchy", "2/4", "List, badges, filters and detail compete.", "Slow scanning", "Use a master-detail grid with vehicle and price dominant."],
        ["Efficiency", "2/4", "Cards repeat dense information.", "Hard comparison", "Standardize differentiating attributes and compare mode."],
        ["Feedback", "2/4", "Selected state does not strongly connect list and detail.", "Context loss", "Strengthen selected state and preserve URL/position."],
        ["Trust", "2/4", "Price, condition, seller and inspection lack structure.", "Risky decision", "Use a normalized trust block and attainable total price."],
        ["Responsive", "1/4", "Three simultaneous zones cannot fit mobile.", "Critical reflow", "List then detail; preserve position and filters on return."],
    ],
    4: [
        ["Hierarchy", "1/4", "Permanent filters consume substantial width.", "Vehicles are obscured", "Use a collapsible drawer and full-width results."],
        ["Efficiency", "2/4", "Many criteria appear simultaneously.", "Cognitive load", "Prioritize facets and place the rest under More filters."],
        ["Feedback", "1/4", "Applied choices and their effect are weakly summarized.", "Disorientation", "Use applied-filter chips and live result counts."],
        ["Prevention", "2/4", "View changes may risk state loss.", "Repeated work", "Persist filters in the URL and confirm Reset."],
        ["Mobile", "1/4", "Sidebar pattern is not thumb-friendly.", "Blocking interaction", "Use a bottom sheet with Show X vehicles."],
    ],
    5: [
        ["Hierarchy", "2/4", "Badges, actions and metadata are compact.", "CTA confusion", "Use one primary CTA and lead with price/availability."],
        ["Comparison", "2/4", "Attributes vary between rows.", "Uneven evaluation", "Use price, mileage, condition, location and seller consistently."],
        ["Trust", "1/4", "Total price and mandatory fees are not evidenced.", "Legal and trust risk", "Disclose attainable price and conditions before commitment."],
        ["Accessibility", "To test", "Small text and targets are suggested by the capture.", "Reading/action difficulty", "Use 16 px base text and measured target sizes."],
        ["Responsive", "2/4", "Dense horizontal actions will wrap poorly.", "Excessive scrolling", "Use a compact vertical card with one dominant action."],
    ],
    6: [
        ["Hierarchy", "2/4", "Badges compete with gallery images.", "Fragmented reading", "Prioritize image, price and model."],
        ["Consistency", "2/4", "List and gallery state continuity is not evident.", "Lost filters", "Share data, sort and persistent selection."],
        ["Empty state", "1/4", "Live Auctions provides little explanation.", "Looks broken", "Add calendar, rules, alerts and upcoming vehicles."],
        ["Trust", "1/4", "Deposit, increment and server time are absent.", "Uncertain commitment", "Keep auction rules and server time visible."],
        ["Mobile", "2/4", "Live bidding is latency-sensitive.", "Bid error", "Use strong confirmation and communicate latency."],
    ],
    7: [
        ["Value", "1/4", "Login-first message provides no seller benefit.", "Abandonment", "Explain VIN, photos, timing and quality review."],
        ["Hierarchy", "2/4", "Signup and Login have no recommended path.", "Unnecessary choice", "Lead with Start a listing and contextual authentication."],
        ["Help", "1/4", "Prerequisites are not previewed.", "Post-account surprise", "Show checklist and estimated time first."],
        ["Accessibility", "To test", "Focus and authentication alternatives are unknown.", "Potential barriers", "Meet WCAG 3.3.8 and expose clear messages."],
        ["Trust", "1/4", "VIN/photo usage is unexplained.", "Data reluctance", "Provide privacy information at collection."],
    ],
    8: [
        ["Clarity", "2/4", "Sparse screen provides limited context.", "Weak next action", "Add purpose, progress and resume-draft option."],
        ["Feedback", "1/4", "Save and progress status are absent.", "Uncertainty", "Show timestamped autosave and draft state."],
        ["Prevention", "2/4", "Required data is not previewed.", "Late errors", "Use progressive validation and examples."],
        ["Consistency", "2/4", "Seller surface feels disconnected from Buy.", "Fragmented product", "Apply one shared design system."],
        ["Mobile", "2/4", "Controls still require validation.", "Target uncertainty", "Use one intention per page and full-width actions."],
    ],
    9: [
        ["Status", "1/4", "Pending appears without explanation.", "More support contacts", "Show cause, owner, date and estimated delay."],
        ["Action", "1/4", "No clear next action.", "Blocked seller", "Offer correct, complete, contact or wait."],
        ["Table", "2/4", "Columns lack hierarchy.", "Slow scanning", "Keep essential columns and row-end actions."],
        ["Trust", "2/4", "Moderation process is unexplained.", "Loss of control", "Show moderation timeline and audit history."],
        ["Responsive", "1/4", "Wide table and empty space.", "Two-dimensional scroll", "Use mobile cards or priority columns."],
    ],
    10: [
        ["Clarity", "1/4", "Finished may combine sold, expired and rejected.", "Ambiguous history", "Use distinct statuses and definitions."],
        ["Efficiency", "2/4", "Reuse and filtering tools are weak.", "Repeated work", "Add filters, controlled duplicate and export."],
        ["Feedback", "1/4", "Outcome and reason are not prominent.", "No learning", "Show outcome, date, price and reason."],
        ["Analytics", "1/4", "Performance and leads are disconnected.", "Lost value", "Show views, contacts, time and conversion."],
        ["Responsive", "2/4", "Simple but low-value empty state.", "Low mobile utility", "Use outcome cards with contextual action."],
    ],
    11: [
        ["Search", "2/4", "Controls are present but unclear.", "Long retrieval time", "Search VIN, stock number, make and status."],
        ["Hierarchy", "2/4", "Results use space inefficiently.", "Reduced productivity", "Use adaptive density and configurable columns."],
        ["Actions", "1/4", "Primary management actions are weak.", "Blocked operations", "Expose edit, publish, pause, duplicate and archive."],
        ["Prevention", "1/4", "Risky actions are not visibly separated.", "Operational error", "Separate routine and destructive actions."],
        ["Accessibility", "To test", "Table/menu keyboard behaviour is unknown.", "Complex navigation", "Use headers, accessible names and managed focus."],
    ],
    12: [
        ["Architecture", "2/4", "Statuses are distributed across weakly differentiated tabs.", "Heavy mental model", "Expose lifecycle and status filters."],
        ["Consistency", "2/4", "Counters and labels vary.", "Reduced confidence", "Use one taxonomy and transition rules."],
        ["Feedback", "1/4", "Status changes are not explained.", "Errors and support", "Show useful toast, history and prior state."],
        ["Prevention", "2/4", "Transitions may be irreversible.", "Lost sale", "Confirm by risk and allow undo where possible."],
        ["Mobile", "1/4", "Tabs and tables can overflow.", "Fragile reflow", "Use status filters and responsive cards."],
    ],
    13: [
        ["Orientation", "2/4", "Open menu does not emphasize role or context.", "Wrong destination", "Show person, organization, role and next task."],
        ["Architecture", "2/4", "Account and product areas appear separated.", "Fragmented navigation", "Use one authenticated application shell."],
        ["Security", "2/4", "Permissions are not visible.", "Role mistakes", "Expose roles and condition actions."],
        ["Feedback", "2/4", "Successful login lacks operational summary.", "Lost time", "Land on a role-personalized dashboard."],
        ["Accessibility", "To test", "Menu focus behaviour is unknown.", "Possible keyboard trap", "Set initial focus, Escape and return focus."],
    ],
    14: [
        ["Clarity", "2/4", "Many fields appear in one large modal.", "High effort", "Lead with VIN and confirm decoded data."],
        ["Progress", "2/4", "Complete/incomplete states are weak.", "Omissions", "Use a stepper with status and errors."],
        ["Prevention", "2/4", "Make/Model/Trim dependencies can conflict.", "Bad data", "Validate dependencies and allow correction."],
        ["Feedback", "1/4", "Save and validation status are weak.", "Perceived data loss", "Show autosave time and adjacent errors."],
        ["Mobile", "1/4", "Dense modal and columns.", "Difficult entry", "Use a dedicated page with short sections."],
    ],
    15: [
        ["Hierarchy", "2/4", "Palettes and many options dominate.", "Slow decision", "Prioritize common/recent values, then search."],
        ["Accessibility", "1/4", "Colour may carry meaning alone.", "Exclusion", "Add text names, selection marks and contrast."],
        ["Consistency", "2/4", "Tyres, colours and attributes use varied patterns.", "Learning burden", "Normalize radio, checkbox and selector patterns."],
        ["Prevention", "2/4", "Technical values lack help.", "Incorrect entry", "Add examples, units and Unknown."],
        ["Mobile", "1/4", "Large palette and many options.", "Excessive scrolling", "Use search, accordions and persistent summary."],
    ],
    16: [
        ["Cognitive load", "1/4", "Long ungrouped tag list.", "Fatigue and omissions", "Group, search and suggest from VIN."],
        ["Hierarchy", "2/4", "All equipment has equal weight.", "Slow scanning", "Lead with key equipment and collapse advanced items."],
        ["Feedback", "2/4", "Selection is visible but overall summary is weak.", "Doubt", "Show count and editable summary."],
        ["Accessibility", "1/4", "Many small toggle targets are likely.", "Difficult interaction", "Meet target size, focus and textual-state rules."],
        ["Performance", "To test", "Large interactive component set.", "Mobile latency", "Use progressive rendering where needed."],
    ],
    17: [
        ["Guidance", "2/4", "Angles and damage capture lack a teaching sequence.", "Incomplete photos", "Provide an illustrated shot checklist."],
        ["Feedback", "1/4", "Upload quality, progress and errors are weak.", "Hard recovery", "Show per-file progress, retry and state."],
        ["Prevention", "2/4", "Blur, duplicates and formats may fail late.", "Late rejection", "Run local checks and explain assisted analysis."],
        ["Privacy", "1/4", "Use, retention and metadata are unexplained.", "Trust risk", "Disclose consent, EXIF removal and retention."],
        ["Mobile", "2/4", "Capture is natural but networks vary.", "Mobile failure", "Compress, resume and support low bandwidth."],
    ],
    18: [
        ["Clarity", "2/4", "Final description lacks a complete review summary.", "Remaining errors", "Add a review page linking to corrections."],
        ["AI assistance", "To govern", "Fast generation lacks visible safeguards.", "False content", "Use factual drafts, field sources and human approval."],
        ["Prevention", "1/4", "Publishing may be irreversible.", "Incorrect listing", "Preview, attest and confirm."],
        ["Feedback", "1/4", "Post-submit status and timing are unclear.", "Uncertainty", "Provide receipt, moderation timeline and notification."],
        ["Trust", "2/4", "Content responsibility is not explicit.", "Dispute risk", "Keep version history and change ownership."],
    ],
}

for title, page, _, _ in SCREEN_AUDIT_EN:
    MATRIX_DATA["EN"][title] = (
        ["Criterion", "Score", "Evidence", "Impact", "Recommendation"],
        EN_AUDIT_POINTS[page],
    )

SCREEN_IMAGE_BY_TITLE = {
    title: Path(f"/tmp/autocommerce-pdf-pages/page-{page:02d}.jpg")
    for title, page, _, _ in SCREEN_AUDIT_FR + SCREEN_AUDIT_EN
}

IMAGE_BY_PAGE = {
    3: Path("/tmp/autocommerce-pdf-pages/page-02.jpg"),
    4: Path("/tmp/autocommerce-pdf-pages/page-04.jpg"),
    5: Path("/tmp/autocommerce-pdf-pages/page-06.jpg"),
    6: Path("/tmp/autocommerce-pdf-pages/page-06.jpg"),
    7: Path("/tmp/autocommerce-pdf-pages/page-09.jpg"),
    8: Path("/tmp/autocommerce-pdf-pages/page-14.jpg"),
    9: Path("/tmp/autocommerce-pdf-pages/page-13.jpg"),
    10: Path("/tmp/autocommerce-pdf-pages/page-18.jpg"),
}


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def configure_document(document: Document, language: dict) -> None:
    section = document.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.68)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    for style_name, size, colour in (
        ("Normal", 9.3, CHARCOAL),
        ("Title", 27, CHARCOAL),
        ("Heading 1", 20, CHARCOAL),
        ("Heading 2", 12.5, GREEN),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(colour)
    document.styles["Title"].font.bold = True
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 2"].font.bold = True
    document.styles["Normal"].paragraph_format.space_after = Pt(4)
    document.styles["Normal"].paragraph_format.line_spacing = 1.05

    footer = section.footer.paragraphs[0]
    footer.text = language["confidential"]
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string(CHARCOAL)
    page = section.footer.add_paragraph()
    add_page_number(page)


def add_brand_bar(document: Document, label: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.25)
    logo = ASSETS / "logo.png"
    if logo.exists():
        table.cell(0, 0).paragraphs[0].add_run().add_picture(str(logo), width=Inches(1.05))
    table.cell(0, 1).text = label
    for cell in table.rows[0].cells:
        set_cell_shading(cell, CHARCOAL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(WHITE)


def add_cover(document: Document, language: dict) -> None:
    add_brand_bar(document, language["subtitle"])
    document.add_paragraph()
    document.add_paragraph()
    title = document.add_paragraph()
    title_run = title.add_run(language["title"])
    title_run.font.name = "Arial"
    title_run.font.size = Pt(27)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string(CHARCOAL)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(80)
    title.paragraph_format.space_after = Pt(18)
    rule = document.add_paragraph("━" * 42)
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.runs[0].font.color.rgb = RGBColor.from_string(RED)
    rule.runs[0].font.size = Pt(10)
    document.add_paragraph()
    signature = document.add_paragraph("\n".join(SIGNATURE))
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, run in enumerate(signature.runs):
        run.font.name = "Arial"
        run.font.size = Pt(10 if index else 11)
        run.font.bold = True
    date_line = document.add_paragraph("JULY 2026")
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.runs[0].font.color.rgb = RGBColor.from_string(GREEN)
    date_line.runs[0].font.bold = True


def add_toc(document: Document, language: dict) -> None:
    entries = list(enumerate(language["pages"], start=1))
    chunks = (entries[:21], entries[21:])
    headers = ("SECTION", "PAGE")
    for chunk_index, chunk in enumerate(chunks):
        if chunk_index:
            document.add_page_break()
        continuation = " — SUITE" if language["code"] == "FR" else " — CONTINUED"
        add_brand_bar(
            document,
            language["toc"] + (continuation if chunk_index else ""),
        )
        heading = document.add_paragraph()
        heading_run = heading.add_run(
            language["toc"] + (continuation.lower() if chunk_index else "")
        )
        heading_run.font.name = "Arial"
        heading_run.font.size = Pt(24 if chunk_index else 27)
        heading_run.font.bold = True
        heading_run.font.color.rgb = RGBColor.from_string(CHARCOAL)
        heading.paragraph_format.space_before = Pt(18)
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.text = header
            set_cell_shading(cell, CHARCOAL)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.font.size = Pt(7.5)
        set_repeat_table_header(table.rows[0])
        for content_index, (title, _) in chunk:
            page_number = content_index + 3
            cells = table.add_row().cells
            cells[0].text = f"{content_index:02d}  {title}"
            cells[1].text = str(page_number)
            for cell in cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(7.5)
            if content_index % 2 == 0:
                set_cell_shading(cells[0], LIGHT)
                set_cell_shading(cells[1], LIGHT)


def add_evidence_image(document: Document, page_number: int, title: str) -> None:
    image = SCREEN_IMAGE_BY_TITLE.get(title)
    if not image or not image.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = 2.45 if title in SCREEN_IMAGE_BY_TITLE else 3.35
    paragraph.add_run().add_picture(str(image), width=Inches(width))
    caption = document.add_paragraph(
        f"Source evidence · supplied AutoCommerce capture · {image.stem}"
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(7.5)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_analysis_page(
    document: Document,
    language: dict,
    content_index: int,
    title: str,
    sections: list[tuple[str, str]],
) -> None:
    page_number = content_index + 3
    add_brand_bar(document, f"{content_index + 1:02d} · {title.upper()}")
    heading = document.add_paragraph(title, style="Heading 1")
    heading.paragraph_format.space_before = Pt(12)
    for section_title, body in sections:
        subheading = document.add_paragraph(section_title, style="Heading 2")
        subheading.paragraph_format.space_before = Pt(6)
        paragraph = document.add_paragraph(body)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    matrix = MATRIX_DATA.get(language["code"], {}).get(title)
    if matrix:
        headers, rows = matrix
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, header in enumerate(headers):
            table.cell(0, idx).text = header
            set_cell_shading(table.cell(0, idx), CHARCOAL)
            for run in table.cell(0, idx).paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(WHITE)
                run.font.size = Pt(7.5)
        set_repeat_table_header(table.rows[0])
        for row_index, values in enumerate(rows):
            cells = table.add_row().cells
            for idx, value in enumerate(values):
                cells[idx].text = value
                cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for run in cells[idx].paragraphs[0].runs:
                    run.font.size = Pt(7.2)
            if row_index % 2:
                for cell in cells:
                    set_cell_shading(cell, LIGHT)

    if title in SCREEN_IMAGE_BY_TITLE:
        add_evidence_image(document, page_number, title)

    if content_index == len(language["pages"]) - 1:
        document.add_paragraph("Sources", style="Heading 2")
        for source in language["sources"]:
            document.add_paragraph(source, style="List Bullet")


def build(language: dict) -> Path:
    document = Document()
    configure_document(document, language)
    add_cover(document, language)
    document.add_page_break()
    add_toc(document, language)
    for index, (title, sections) in enumerate(language["pages"]):
        document.add_page_break()
        add_analysis_page(document, language, index, title, sections)

    OUTPUT.mkdir(parents=True, exist_ok=True)
    output_path = OUTPUT / f"autocommerce-document-01-{language['code'].lower()}.docx"
    document.save(output_path)
    return output_path


if __name__ == "__main__":
    for payload in (FR, EN):
        print(build(payload))
