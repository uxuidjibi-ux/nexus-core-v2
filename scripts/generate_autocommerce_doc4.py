from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "Autocommerce" / "document-04"

BLACK, RED, GREEN = "171717", "E00810", "104830"
INK, GREY, LIGHT, WHITE = "292D32", "677078", "F2F4F4", "FFFFFF"
BLUE, AMBER = "2364AA", "D89216"
SIGNATURE = (
    "PREPARED BY DJIGO DJIBI",
    "CX Consultant | Strategic Product & UX/UI Designer | AI Front-End Developer",
)

SOURCES = [
    ("Figma REST API — File endpoints", "https://developers.figma.com/docs/rest-api/file-endpoints/"),
    ("Figma REST API — Variables endpoints", "https://developers.figma.com/docs/rest-api/variables-endpoints/"),
    ("Make — Webhooks", "https://help.make.com/webhooks"),
    ("Make — Error handlers", "https://help.make.com/error-handlers"),
    ("Make — Retry error handler", "https://help.make.com/retry-error-handler"),
    ("WordPress — REST API handbook", "https://developer.wordpress.org/rest-api/"),
    ("WordPress — REST authentication", "https://developer.wordpress.org/rest-api/using-the-rest-api/authentication/"),
    ("W3C — WCAG 2.2", "https://www.w3.org/TR/WCAG22/"),
]

FR = {
    "lang": "FR",
    "title": "AUTOCOMMERCE.CA — DOCUMENT 04\nMVP PRODUCT REQUIREMENTS SPECIFICATION",
    "subtitle": "FIGMA PROTOTYPES · MAKE AUTOMATION · READYAI / CUSTOM AGENTS",
    "toc": "Table des matières",
    "conf": "AUTOCOMMERCE · DOCUMENT 04 · CONFIDENTIEL · FR",
    "pages": [
        ("1. Périmètre du MVP & arbitrage 80/20",
         "Le MVP valide une promesse unique : permettre à un acheteur de trouver et comprendre un véhicule rapidement, et à un commerçant de publier une annonce fiable en quelques minutes grâce à une assistance contrôlée.",
         "scope"),
        ("2. Prototypes Figma & parcours prioritaires",
         "Le prototype Figma doit démontrer cinq parcours critiques sur ordinateur et mobile, avec une architecture de composants réutilisable et suffisamment documentée pour une future implémentation.",
         "figma_flows"),
        ("2. Design tokens, composants & états",
         "Le système visuel conserve le logo et les couleurs AutoCommerce, mais utilise le blanc et les gris clairs comme surfaces principales. Le noir structure la hiérarchie; rouge et vert sont réservés aux actions et états significatifs.",
         "figma_tokens"),
        ("3. Architecture Make / Zapier",
         "L’automatisation ne remplace pas le backend : elle orchestre l’ingestion, la validation, la génération IA, la publication et les notifications derrière un contrat d’événements versionné.",
         "automation"),
        ("3. Fiabilité, erreurs, logs & sécurité",
         "Chaque scénario est idempotent, observable et récupérable. Une erreur temporaire est réessayée; une donnée invalide est mise en quarantaine; une action irréversible exige une validation ou une clé d’idempotence.",
         "reliability"),
        ("4. ReadyAI & custom agents — contrats",
         "ReadyAI est traité comme un fournisseur interchangeable derrière un adaptateur. Aucun endpoint propriétaire n’est présumé tant que sa documentation et ses accès ne sont pas confirmés.",
         "ai_contract"),
        ("4. Prompts, garde-fous & fallbacks",
         "L’IA transforme uniquement des faits fournis ou vérifiés. Elle peut reformuler et hiérarchiser, jamais inventer l’état, l’historique, le prix, la disponibilité ou une garantie.",
         "ai_guardrails"),
        ("5. Plan de test, recette & KPIs",
         "La recette couvre le parcours complet depuis l’ingestion VIN/médias jusqu’à l’affichage public, avec des cas passants, limites, erreurs, performance, accessibilité et reprise.",
         "qa"),
        ("6. Roadmap V2/V3 & scalabilité",
         "Le MVP reste volontairement étroit. La V2 industrialise les opérations et la confiance; la V3 ouvre la transaction, l’intelligence réseau et l’offre Enterprise.",
         "roadmap"),
        ("Annexe opérationnelle — décisions & sources",
         "Ce cahier des charges est prêt pour prototypage. Les valeurs marquées « cible » sont des seuils produit à valider; les capacités fournisseurs doivent être vérifiées au moment de l’intégration.",
         "sources"),
    ],
}

EN = {
    "lang": "EN",
    "title": "AUTOCOMMERCE.CA — DOCUMENT 04\nMVP PRODUCT REQUIREMENTS SPECIFICATION",
    "subtitle": "FIGMA PROTOTYPES · MAKE AUTOMATION · READYAI / CUSTOM AGENTS",
    "toc": "Table of contents",
    "conf": "AUTOCOMMERCE · DOCUMENT 04 · CONFIDENTIAL · EN",
    "pages": [
        ("1. MVP scope & 80/20 arbitration",
         "The MVP validates one promise: help buyers find and understand a vehicle quickly, while enabling merchants to publish a reliable listing in minutes through controlled assistance.",
         "scope"),
        ("2. Figma prototypes & priority journeys",
         "The Figma prototype must demonstrate five critical desktop and mobile journeys, supported by reusable components and documentation ready for future implementation.",
         "figma_flows"),
        ("2. Design tokens, components & states",
         "The visual system preserves the AutoCommerce logo and colours while using white and light grey as primary surfaces. Black structures hierarchy; red and green are reserved for meaningful actions and states.",
         "figma_tokens"),
        ("3. Make / Zapier architecture",
         "Automation does not replace the backend. It orchestrates ingestion, validation, AI generation, publishing and notifications behind a versioned event contract.",
         "automation"),
        ("3. Reliability, errors, logs & security",
         "Every scenario is idempotent, observable and recoverable. Temporary failures are retried; invalid data is quarantined; irreversible actions require approval or an idempotency key.",
         "reliability"),
        ("4. ReadyAI & custom agents — contracts",
         "ReadyAI is treated as an interchangeable provider behind an adapter. No proprietary endpoint is assumed until documentation and access are confirmed.",
         "ai_contract"),
        ("4. Prompts, guardrails & fallbacks",
         "AI transforms only supplied or verified facts. It may rewrite and prioritize, but never invent condition, history, price, availability or a warranty.",
         "ai_guardrails"),
        ("5. Test plan, acceptance & KPIs",
         "Acceptance covers the complete journey from VIN/media ingestion to public display, including happy paths, edge cases, errors, performance, accessibility and recovery.",
         "qa"),
        ("6. V2/V3 roadmap & scalability",
         "The MVP remains intentionally narrow. V2 industrializes operations and trust; V3 unlocks transaction, network intelligence and the Enterprise offering.",
         "roadmap"),
        ("Operational appendix — decisions & sources",
         "This specification is ready for prototyping. Values marked “target” are product thresholds to validate; provider capabilities must be confirmed at integration time.",
         "sources"),
    ],
}


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    trpr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    trpr.append(tag)


def field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, sep, end])


def setup(doc, conf):
    sec = doc.sections[0]
    sec.top_margin = Inches(.62)
    sec.bottom_margin = Inches(.62)
    sec.left_margin = Inches(.72)
    sec.right_margin = Inches(.72)
    sec.header_distance = Inches(.3)
    sec.footer_distance = Inches(.3)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [("Heading 1", 20, BLACK), ("Heading 2", 12.5, RED), ("Heading 3", 10.5, GREEN)]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(5)
        st.paragraph_format.space_after = Pt(6)
    header = sec.header.paragraphs[0]
    header.text = conf
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        r.font.name = "Arial"
        r.font.size = Pt(7)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(GREY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("AUTOCOMMERCE  ·  ")
    field(footer, "PAGE")
    for r in footer.runs:
        r.font.name = "Arial"
        r.font.size = Pt(7)
        r.font.color.rgb = RGBColor.from_string(GREY)


def table(doc, rows, widths=None, size=7.1):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.style = "Table Grid"
    t.autofit = False
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, values in enumerate(rows):
        cells = t.add_row().cells
        for j, value in enumerate(values):
            if widths:
                cells[j].width = Inches(widths[j])
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(str(value))
            r.font.name = "Arial"
            r.font.size = Pt(size)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(WHITE)
                shade(cells[j], BLACK if j == 0 else GREEN)
            elif j == 0:
                r.bold = True
                shade(cells[j], LIGHT)
        if i == 0:
            repeat_header(t.rows[-1])
    doc.add_paragraph()


def callout(doc, label, body, color=GREEN):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.cell(0, 0).width = Inches(6.9)
    shade(t.cell(0, 0), LIGHT)
    p = t.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(body)
    doc.add_paragraph()


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(.22)
        p.paragraph_format.first_line_indent = Inches(-.12)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def mono(doc, lines):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t.cell(0, 0).width = Inches(6.9)
    shade(t.cell(0, 0), "202428")
    p = t.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Menlo"
        r.font.size = Pt(7.2)
        r.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph()


def cover(doc, data):
    doc.add_paragraph("\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(55)
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(data["title"])
    r.font.name = "Arial"
    r.font.size = Pt(27)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLACK)
    line = doc.add_paragraph("━" * 34)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.runs[0].font.color.rgb = RGBColor.from_string(RED)
    sub = doc.add_paragraph(data["subtitle"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph("\n\n")
    for i, text in enumerate(SIGNATURE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11 if i == 0 else 9)
        r.bold = i == 0
        r.font.color.rgb = RGBColor.from_string(GREEN if i == 0 else BLACK)


def toc(doc, data):
    doc.add_heading(data["toc"], 0)
    fr = data["lang"] == "FR"
    rows = [
        ["#", "Chapitre" if fr else "Chapter", "Contenu" if fr else "Content", "Page"],
        ["1", "Périmètre MVP" if fr else "MVP scope", "80/20 · MoSCoW · critères d’acceptation" if fr else "80/20 · MoSCoW · acceptance", "3"],
        ["2", "Figma & design system", "Flows · composants · tokens · états" if fr else "Flows · components · tokens · states", "4–5"],
        ["3", "Automation", "Make/Zapier · WordPress · erreurs · logs" if fr else "Make/Zapier · WordPress · errors · logs", "6–7"],
        ["4", "ReadyAI & agents", "Prompts · JSON · garde-fous · fallbacks" if fr else "Prompts · JSON · guardrails · fallbacks", "8–9"],
        ["5", "QA & métriques" if fr else "QA & metrics", "E2E · recette · performance · KPIs", "10"],
        ["6", "Roadmap & scalabilité" if fr else "Roadmap & scalability", "V2 · V3 · architecture cible" if fr else "V2 · V3 · target architecture", "11"],
        ["A", "Annexe" if fr else "Appendix", "Décisions, dépendances et sources" if fr else "Decisions, dependencies and sources", "12"],
    ]
    table(doc, rows, [0.35, 1.8, 3.9, .75], 7.4)
    callout(doc, "Vision", "Simple au premier regard, puissant à la demande. L’utilisateur reste maître des filtres, des données et des contenus générés." if fr else "Simple at first glance, powerful on demand. Users remain in control of filters, data and generated content.")


def heading(doc, index, title, intro):
    p = doc.add_paragraph()
    r = p.add_run(f"{index:02d} · MVP PRODUCT REQUIREMENTS")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(RED)
    doc.add_heading(title, 0)
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(9)
    p.runs[0].font.size = Pt(10.1)
    p.runs[0].font.color.rgb = RGBColor.from_string(GREY)


def content(doc, kind, lang):
    fr = lang == "FR"
    if kind == "scope":
        callout(doc, "Décision MVP" if fr else "MVP decision",
                "Trois preuves à obtenir : rechercher sans friction, publier une annonce assistée fiable, agir depuis un cockpit commerçant." if fr else "Three proofs: frictionless search, reliable assisted listing creation and action from a merchant cockpit.")
        scope_rows = [
            ["MoSCoW", "Fonctions", "Raison / limite"],
            ["MUST", "Recherche, résultats, filtre rétractable, VDP, favoris; VIN, photos, description, prévisualisation; inventaire et prospects", "Couvre la boucle de valeur complète"],
            ["SHOULD", "Comparaison 2–4 véhicules, alertes, rendez-vous, score qualité, import CSV", "Augmente décision et productivité"],
            ["COULD", "Estimation financement, recommandations, multilingue enrichi, landing pages d’offres", "À tester sans dépendance critique"],
            ["WON’T", "Paiement, enchères live, escrow, shipping, app native, pricing dynamique autonome", "Risque/compliance trop élevés pour le MVP"],
        ] if fr else [
            ["MoSCoW", "Capabilities", "Rationale / limit"],
            ["MUST", "Search, results, collapsible filter, VDP, favourites; VIN, photos, description, preview; inventory and leads", "Covers the complete value loop"],
            ["SHOULD", "Compare 2–4 vehicles, alerts, appointments, quality score, CSV import", "Improves decision and productivity"],
            ["COULD", "Financing estimate, recommendations, enhanced multilingual content, offer landing pages", "Test without critical dependency"],
            ["WON’T", "Payment, live auction, escrow, shipping, native app, autonomous dynamic pricing", "Too much risk/compliance for MVP"],
        ]
        table(doc, scope_rows, [1.0, 4.0, 2.0], 6.8)
        acceptance_rows = [
            ["Surface", "Core feature", "Acceptation"],
            ["Search", "3 champs + localisation; filtres en tiroir; chips actifs", "Recherche conservée au changement de vue"],
            ["VDP", "Prix total, historique, état, vendeur, CTA contextuelle", "Données critiques sourcées et datées"],
            ["Listing", "VIN/manual, photo QA, prix, texte assisté, preview", "Aucune publication sans validation humaine"],
        ] if fr else [
            ["Surface", "Core feature", "Acceptance"],
            ["Search", "3 fields + location; filter drawer; applied chips", "Search persists across view changes"],
            ["VDP", "Total price, history, condition, seller, contextual CTA", "Critical data sourced and dated"],
            ["Listing", "VIN/manual, photo QA, price, assisted copy, preview", "No publish without human validation"],
        ]
        table(doc, acceptance_rows, [1.0, 3.5, 2.5], 7.0)
    elif kind == "figma_flows":
        mono(doc, [
            "BUYER  Home → Search → Results ⇄ Filter Drawer → Compare → VDP → Contact/Appointment",
            "SELLER Login → VIN/Manual → Facts → Condition → Guided Photos → Price/AI Copy → Preview",
            "MERCHANT Dashboard → Inventory → Listing QA → Publish → Lead Inbox → Follow-up",
            "MOBILE  Search → Results → Bottom-sheet Filters → VDP → Sticky Primary Action",
        ])
        table(doc, [
            ["Flow", "Frames desktop", "Frames mobile", "Prototype proof"],
            ["Acheter" if fr else "Buy", "1440×1024: home, results, compare, VDP", "390×844: same task", "Filter state persists"],
            ["Vendre" if fr else "Sell", "Wizard + preview + success", "Step-by-step + camera", "Draft autosave"],
            ["Pilotage" if fr else "Operate", "Dashboard, inventory, lead detail", "KPIs + tasks + lead", "Next action visible"],
        ], [1.0, 2.3, 2.1, 1.6], 6.9)
        bullets(doc, [
            "Nommage : 01_Foundations, 02_Components, 03_Patterns, 10_Buyer, 20_Seller, 30_Merchant, 90_Archive." if fr else "Naming: 01_Foundations, 02_Components, 03_Patterns, 10_Buyer, 20_Seller, 30_Merchant, 90_Archive.",
            "Auto Layout sur chaque composant; contraintes de redimensionnement; propriétés documentées." if fr else "Auto Layout on every component; resizing constraints; documented properties.",
            "Prototype clavier et mobile; focus, retours, transitions ≤200 ms hors chargement réel." if fr else "Keyboard and mobile prototype; focus, feedback, transitions ≤200 ms except real loading.",
        ])
        callout(doc, "Definition of Done", "5 parcours cliquables, desktop/mobile, sans écran mort; tous les états clés reliés." if fr else "5 clickable desktop/mobile journeys, no dead ends; all key states linked.", BLUE)
    elif kind == "figma_tokens":
        table(doc, [
            ["Collection", "Token examples", "Rule"],
            ["Color", "surface/default #FFFFFF · ink #171717 · action/red #E00810 · trust/green #104830", "Semantic names, not component names"],
            ["Typography", "display/32 · h1/24 · h2/18 · body/16 · meta/13", "Minimum 16 px for primary mobile body"],
            ["Spacing", "space/1 4 · /2 8 · /3 12 · /4 16 · /6 24 · /8 32", "4 px base; no arbitrary values"],
            ["Radius", "sm 6 · md 10 · lg 16 · pill 999", "Cards md; chips pill"],
            ["Grid", "Desktop 12×72/24 · Tablet 8×64/20 · Mobile 4×78/16", "Fluid gutters at breakpoints"],
        ], [1.1, 3.7, 2.2], 6.8)
        table(doc, [
            ["Component", "Variants", "Required states"],
            ["Button", "Primary, secondary, tertiary, destructive, icon", "Default · hover · focus · pressed · disabled · loading"],
            ["Filter", "Select, range, checkbox, chip, drawer", "Applied · invalid · empty · reset"],
            ["Vehicle card", "Grid, list, compact", "Default · saved · compared · unavailable · skeleton"],
            ["Form field", "Text, number, VIN, upload", "Empty · valid · error · warning · readonly"],
            ["Feedback", "Toast, inline, banner, modal", "Info · success · warning · error"],
        ], [1.25, 2.45, 3.3], 6.9)
        callout(doc, "Figma API", "GET /v1/files/:key lit l’arbre, composants et styles. L’écriture Variables REST est Enterprise et exige l’accès édition; prévoir un export JSON de secours." if fr else "GET /v1/files/:key reads the tree, components and styles. REST Variables writing is Enterprise-only and requires edit access; provide a JSON export fallback.", AMBER)
    elif kind == "automation":
        mono(doc, [
            "[Webhook inventory/listing.v1]",
            "       ↓ verify signature + schema + idempotency_key",
            "[Parse/Normalize] → [Router: create | update | reject]",
            "       ↓                    ↘ [Quarantine + alert]",
            "[DB upsert] → [Media QA] → [AI adapter] → [Human approval]",
            "       ↓                                       ↓",
            "[WordPress/Headless draft] ← [Publish command] → [Event + notification]",
        ])
        table(doc, [
            ["Stage", "Make module / pattern", "Contract"],
            ["Trigger", "Custom Webhook or scheduled inventory poll", "event_id, type, occurred_at, schema_version"],
            ["Transform", "JSON parse, iterator, aggregator, formatter", "Canonical vehicle.v1 object"],
            ["Decision", "Router + filters", "create/update/reject; no silent drop"],
            ["Destination", "HTTP → backend/WordPress; DB/data store", "Draft first; publish only after approval"],
            ["Response", "Webhook response last", "202 accepted + correlation_id"],
        ], [1.1, 2.7, 3.2], 6.9)
        callout(doc, "Architecture rule", "Make orchestre; le backend détient les données, les permissions, l’audit et les règles métier." if fr else "Make orchestrates; the backend owns data, permissions, audit and business rules.", BLUE)
    elif kind == "reliability":
        table(doc, [
            ["Failure", "Policy", "Retry / exit", "Evidence"],
            ["429 / connection", "Retry exponential + jitter", "3 attempts; respect Retry-After", "attempt_count + duration"],
            ["5xx / timeout", "Retry then incomplete execution", "1m, 5m, 15m", "endpoint + correlation_id"],
            ["400 / schema", "No retry; quarantine", "Manual correction", "field_errors"],
            ["401 / 403", "Stop and alert security owner", "Credential rotation", "secret_ref, never secret"],
            ["Duplicate", "Idempotent return", "200 existing result", "idempotency_key"],
            ["Partial publish", "Compensate / rollback draft", "Manual review", "before/after snapshot"],
        ], [1.15, 2.1, 1.85, 1.9], 6.5)
        bullets(doc, [
            "Journal minimal : event_id, correlation_id, scenario_version, module, status, latency_ms, attempt, sanitized_error." if fr else "Minimum log: event_id, correlation_id, scenario_version, module, status, latency_ms, attempt, sanitized_error.",
            "Secrets dans le coffre Make/Hostinger ou variables d’environnement; jamais dans Figma, prompt, log ou payload." if fr else "Secrets in Make/Hostinger vault or environment variables; never in Figma, prompts, logs or payloads.",
            "WordPress : Application Password distinct, HTTPS, utilisateur à privilèges minimaux, révocation documentée." if fr else "WordPress: dedicated Application Password, HTTPS, least-privileged user, documented revocation.",
            "Conservation cible : logs métier 90 jours; erreurs techniques 30 jours; médias selon politique juridique." if fr else "Target retention: business logs 90 days; technical errors 30 days; media per legal policy.",
        ])
        callout(doc, "SLO MVP", "≥99 % des événements valides traités sans intervention; zéro publication dupliquée; 100 % des échecs traçables." if fr else "≥99% of valid events processed without intervention; zero duplicate publications; 100% of failures traceable.")
    elif kind == "ai_contract":
        mono(doc, [
            '{ "schema_version":"vehicle.v1", "request_id":"uuid",',
            '  "locale":"fr-CA", "task":"listing_description",',
            '  "vehicle":{"vin":"…","year":2024,"make":"…","model":"…",',
            '    "trim":"…","mileage_km":18500,"price_cad":38990,',
            '    "features":["AWD"],"condition_notes":["…"]},',
            '  "constraints":{"max_words":140,"claims_policy":"verified_only"} }',
            "",
            '{ "status":"draft","headline":"…","summary":"…",',
            '  "bullets":["…"],"facts_used":["year","make","model"],',
            '  "warnings":[],"confidence":0.92,"model_trace":"provider/model/version" }',
        ])
        table(doc, [
            ["Agent", "Responsibility", "Input → output"],
            ["LISTING_WRITER", "Fact-based vehicle copy", "vehicle.v1 → listing_copy.v1"],
            ["MEDIA_QA", "Blur, duplicates, missing angles", "media_manifest.v1 → media_report.v1"],
            ["OFFER_PAGE", "Landing-page blocks from approved offer", "offer.v1 → page_blocks.v1"],
            ["COMPLIANCE_REVIEWER", "Flag unsupported claims and sensitive data", "draft + facts → review.v1"],
        ], [1.4, 3.0, 2.6], 6.9)
        callout(doc, "Provider adapter", "readyai.generate(task, payload, policy) → normalized_result. Le même contrat permet un fallback vers un autre moteur sans modifier le produit." if fr else "readyai.generate(task, payload, policy) → normalized_result. The same contract enables another engine without changing the product.", BLUE)
    elif kind == "ai_guardrails":
        mono(doc, [
            "SYSTEM",
            "You are AutoCommerce Listing Writer. Use only provided verified fields.",
            "Never infer accident history, warranty, ownership, availability or price.",
            "Return JSON matching listing_copy.v1. If a critical fact is missing,",
            "add it to warnings and omit the claim. French-Canadian, factual, concise.",
        ])
        table(doc, [
            ["Control", "Rule", "Fallback"],
            ["Schema", "Strict JSON validation; reject extra critical fields", "Template generated from verified facts"],
            ["Moderation", "Block hate, sexual content, personal data, fraud patterns", "Quarantine + human review"],
            ["Grounding", "Every sentence maps to facts_used", "Remove unsupported sentence"],
            ["Latency", "Soft 8 s; hard 20 s", "Async job + status; deterministic template"],
            ["Availability", "Circuit breaker after repeated failures", "Queue + manual copy editor"],
            ["Human control", "Preview + edit + explicit approve", "Save as draft only"],
        ], [1.2, 3.7, 2.1], 6.7)
        bullets(doc, [
            "Prompts versionnés dans Git; changement soumis à revue et jeu de tests de non-régression." if fr else "Prompts versioned in Git; changes require review and regression tests.",
            "PII minimisée; VIN traité selon politique interne; aucune donnée d’authentification envoyée au modèle." if fr else "Minimize PII; handle VIN under internal policy; never send credentials to the model.",
            "Mesurer taux de JSON valide, corrections humaines, claims bloqués, latence et coût par génération." if fr else "Measure valid JSON rate, human edits, blocked claims, latency and cost per generation.",
        ])
    elif kind == "qa":
        table(doc, [
            ["ID", "Scenario", "Expected", "Threshold"],
            ["E2E-01", "VIN valid → media → AI → approval → publish", "One correct public listing", "100% critical fields"],
            ["E2E-02", "Manual entry without VIN", "Draft allowed; warning visible", "No invented fact"],
            ["E2E-03", "Duplicate webhook", "No duplicate listing", "0 duplicates"],
            ["E2E-04", "AI timeout / invalid JSON", "Fallback draft + alert", "No user data loss"],
            ["E2E-05", "WordPress 429/5xx", "Retry then incomplete execution", "Traceable recovery"],
            ["E2E-06", "Mobile filters + VDP", "State retained; keyboard/focus works", "WCAG 2.2 AA target"],
            ["LOAD-01", "50 concurrent listing events", "Queue stable; no corruption", "<1% failed after retry"],
        ], [.65, 2.55, 2.8, 1.0], 6.3)
        table(doc, [
            ["KPI", "MVP target", "Instrumentation"],
            ["Search result LCP", "≤2.5 s p75 mobile", "RUM / Web Vitals"],
            ["Search → VDP", "Baseline + weekly trend", "product analytics event"],
            ["Lead form conversion", "≥8% qualified VDP sessions", "form_success / VDP"],
            ["Listing creation time", "≤8 min median", "wizard start → approved draft"],
            ["AI draft latency", "≤10 s p95; fallback at 20 s", "provider trace"],
            ["AI acceptance", "≥70% with light edit", "edit distance + approval"],
            ["Automation success", "≥99% valid events", "scenario logs"],
        ], [2.1, 2.1, 2.8], 6.7)
        callout(doc, "Go / No-Go", "Aucun défaut critique sécurité/données; 0 blocage sur les 5 parcours; ≥95 % des cas prioritaires réussis." if fr else "No critical security/data defect; zero blockers across 5 journeys; ≥95% priority test pass rate.", RED)
    elif kind == "roadmap":
        table(doc, [
            ["Phase", "Product scope", "Technical evolution", "Exit gate"],
            ["MVP · 0–3 mois", "Search, VDP, assisted listing, merchant cockpit", "Modular monolith + Make + provider adapter", "5 journeys validated"],
            ["V2 · 3–9 mois", "Comparison, alerts, appointments, teams, quality/pricing insights", "Queue, object storage, search index, RBAC, observability", "Repeatable merchant adoption"],
            ["V3 · 9–18 mois", "Offers, auction, payments/escrow partner, logistics, Enterprise API", "Event bus, services by load/domain, warehouse, DR", "Compliance + unit economics"],
        ], [1.1, 2.25, 2.5, 1.15], 6.5)
        mono(doc, [
            "WEB / MOBILE WEB",
            "      ↓ API Gateway + Auth",
            "[Product API] — [Inventory DB] — [Object Storage]",
            "      ↓ events          ↘ [Search Index]",
            "[Queue/Workers] → [AI Adapter] → [Make connectors]",
            "      ↓                         ↘ WordPress / CRM / Notifications",
            "[Audit + Metrics + Traces + Alerts]",
        ])
        bullets(doc, [
            "Commencer par un monolithe modulaire; extraire un service uniquement avec preuve de charge ou d’autonomie." if fr else "Start with a modular monolith; extract a service only with load or autonomy evidence.",
            "CDN pour médias, images responsives, cache de recherche, pagination cursor, traitements lourds asynchrones." if fr else "Use CDN media, responsive images, search caching, cursor pagination and async heavy processing.",
            "Backups chiffrés, restauration testée, environnements séparés, migrations réversibles, feature flags." if fr else "Encrypted backups, tested restore, separate environments, reversible migrations and feature flags.",
        ])
    elif kind == "sources":
        table(doc, [
            ["Decision", "Owner", "Before build"],
            ["ReadyAI provider and API", "CYPHER", "Obtain official docs, auth, limits, DPA and sandbox"],
            ["Figma plan / Variables API", "PIXEL", "Confirm Enterprise rights or use JSON/token plugin fallback"],
            ["CMS role", "WEAVER", "Choose WordPress as presentation layer vs source of truth"],
            ["Vehicle/history data", "AURA + Legal", "Confirm VIN/CARFAX licensing and consent"],
            ["AI policy", "CYPHER + Product", "Approve claim, privacy, retention and human-review rules"],
        ], [2.0, 1.35, 3.65], 6.8)
        rows = [["#", "Source", "URL"]]
        for i, (name, url) in enumerate(SOURCES, 1):
            rows.append([str(i), name, url])
        table(doc, rows, [.35, 3.1, 3.55], 6.0)
        callout(doc, "Final gate", "Prototype review → technical spike → security review → user test → backlog signed off → MVP build." if fr else "Prototype review → technical spike → security review → user test → signed-off backlog → MVP build.", BLUE)


def make(data):
    doc = Document()
    setup(doc, data["conf"])
    cover(doc, data)
    doc.add_page_break()
    toc(doc, data)
    for i, (title, intro, kind) in enumerate(data["pages"], 1):
        doc.add_page_break()
        heading(doc, i, title, intro)
        content(doc, kind, data["lang"])
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / f"autocommerce-document-04-{data['lang'].lower()}.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    for payload in (FR, EN):
        print(make(payload))
