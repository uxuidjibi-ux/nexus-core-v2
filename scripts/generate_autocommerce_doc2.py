from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "artifacts" / "Autocommerce" / "document-02"
CAPTURES = {
    "cars": Path("/private/tmp/cars-home.png"),
    "cargurus": Path("/private/tmp/cargurus-home.png"),
    "autotrader": Path("/private/tmp/autotrader-home.png"),
    "autowini": Path("/private/tmp/autowini-home.png"),
    "cars_results": Path("/private/tmp/cars-results.png"),
    "autotrader_results": Path("/private/tmp/autotrader-results.png"),
}

BLACK, RED, GREEN = "171717", "E00810", "104830"
INK, GREY, LIGHT, WHITE = "2D3033", "697077", "F3F5F5", "FFFFFF"
BLUE, AMBER = "2364AA", "E9A23B"
SIGNATURE = (
    "PREPARED BY",
    "DJIGO DJIBI",
    "CX Consultant | Strategic Product & UX/UI Designer | AI Front-End Developer",
)

SOURCES = [
    ("Cars.ca — marketplace and observed search journey", "https://www.cars.ca/fr"),
    ("CarGurus Canada — marketplace", "https://www.cargurus.ca/"),
    ("CarGurus — mission, deal ratings and seller value", "https://www.cargurus.ca/about/cargurus"),
    ("CarGurus — dealer review integrity", "https://www.cargurus.ca/about/dealer-reviews"),
    ("CarGurus — online purchase and financing journey", "https://www.cargurus.ca/shop/start-your-purchase-online"),
    ("AutoHebdo — marketplace", "https://www.autohebdo.net/"),
    ("AutoHebdo — dealer offer and private listing journey", "https://www.autohebdo.net/Ico/"),
    ("Autowini — purchase, payment and shipping steps", "https://www.autowini.com/en/help/steps"),
    ("Autowini — business membership", "https://m.autowini.com/en/help/business-membership"),
    ("W3C — Web Content Accessibility Guidelines 2.2", "https://www.w3.org/TR/WCAG22/"),
    ("Nielsen Norman Group — 10 usability heuristics", "https://www.nngroup.com/articles/ten-usability-heuristics/"),
    ("Baymard Institute — Product Lists & Filtering UX", "https://baymard.com/research/ecommerce-product-lists"),
]

FR = {
    "lang": "FR",
    "title": "Étude comparative du marché\nBenchmark UX/UI & CX",
    "subtitle": "AUTOCOMMERCE · DOCUMENT 02 · JUILLET 2026",
    "toc": "Table des matières",
    "conf": "AUTOCOMMERCE · CONFIDENTIEL · FR",
    "sections": [
        ("Synthèse exécutive", "Cars.ca offre le socle de découverte le plus sobre; CarGurus transforme les données de prix et de réputation en aide à la décision; AutoHebdo combine une très grande profondeur d’inventaire avec un parcours plus dense; Autowini orchestre la transaction internationale, l’inspection et l’expédition. AutoCommerce ne doit pas juxtaposer ces modèles : il doit réunir leur valeur dans une expérience canadienne plus simple, progressive et pilotable par les commerçants."),
        ("Mandat, périmètre et méthode", "Benchmark expert des parcours publics Acheteur et Vendeur de Cars.ca, CarGurus Canada, AutoHebdo/AutoTrader et Autowini. L’évaluation combine observation d’interface, informations officielles et critères UX/CX. Elle ne constitue ni un test utilisateur, ni un audit technique, ni une certification WCAG."),
        ("Paysage concurrentiel", "Les concurrents occupent quatre territoires distincts : simplicité de découverte, intelligence de marché, puissance de place de marché et commerce transfrontalier. Cet écart crée un espace pour une plateforme qui réduit l’effort de l’acheteur tout en augmentant la productivité du vendeur."),
        ("Cars.ca — profil", "Référence fonctionnelle prioritaire pour AutoCommerce : interface claire, recherche immédiate, bilinguisme canadien, favoris, baisses de prix, recherche par modèle et cartes lisibles. Les résultats observés proposent des filtres compacts, des vues grille/liste, des badges de prix et des rapports CARFAX gratuits. L’expérience est efficace, mais le parcours transactionnel et les outils commerçants sont moins visibles publiquement."),
        ("CarGurus — profil", "La proposition repose sur la confiance et l’aide à la décision : Deal Ratings, détails véhicule, avis concessionnaires, estimation budgétaire, financement amorcé en ligne, prise de rendez-vous et estimation vendeur. Cette intelligence est différenciante; elle exige toutefois des explications transparentes pour éviter que le score ne soit perçu comme une boîte noire."),
        ("AutoHebdo — profil", "AutoHebdo se distingue par la largeur d’inventaire, la profondeur des filtres, les alertes, favoris et deux voies de vente : offre de concessionnaire ou annonce privée gratuite. Cette puissance répond aux experts, mais la densité, les longues descriptions, certaines ruptures de couleur et les interruptions peuvent concurrencer la tâche principale."),
        ("Autowini — profil", "Autowini traite la vente comme une chaîne internationale : réservation, paiement, suivi de commande, contrôle qualité au port, expédition, guides pays et membership B2B. La richesse opérationnelle est forte, mais la page d’accueil très promotionnelle, la navigation dense et certains accès limités augmentent la charge cognitive."),
        ("Comparaison visuelle des accueils", "Les quatre captures montrent des stratégies opposées : Cars.ca et CarGurus concentrent l’attention sur la recherche; AutoHebdo place davantage de navigation, promotions et consentement autour de l’action; Autowini expose simultanément catégories, offres, services logistiques et membership."),
        ("Recherche et découverte", "La meilleure expérience combine une entrée simple — marque, modèle, localisation — avec des chemins alternatifs par budget, carrosserie, usage ou recommandation. AutoCommerce doit conserver une recherche centrale, révéler les critères avancés progressivement et mémoriser les recherches sans exiger un compte prématurément."),
        ("Résultats, filtres et comparaison", "Cars.ca fournit le meilleur point de départ visuel. AutoHebdo démontre la valeur d’un grand vocabulaire de filtres, mais aussi le risque d’un panneau dominant. La cible AutoCommerce est un tiroir rétractable, des pastilles de critères actifs, une comparaison persistante et un tri expliquant son effet."),
        ("Fiche véhicule et confiance", "Une décision automobile exige prix total, historique, état, qualité des photos, vendeur, disponibilité et prochaines étapes. Les forces de CarGurus (valeur et avis), Cars.ca (lisibilité et CARFAX) et Autowini (inspection/logistique) doivent converger dans une fiche AutoCommerce factuelle, explicable et non promotionnelle."),
        ("Expérience vendeur et commerçant", "AutoHebdo simplifie le choix entre rapidité concessionnaire et valeur potentielle de la vente privée; CarGurus fournit estimation et mise en relation; Autowini sert les acheteurs professionnels récurrents. AutoCommerce peut gagner avec un cockpit commerçant : inventaire, score qualité, prospects, offres, rendez-vous, tâches, équipe et performance."),
        ("Compte, fidélisation et service", "Favoris, alertes de prix et recherches enregistrées soutiennent le retour. L’avantage décisif vient d’une continuité complète : activité récente, messages, rendez-vous, documents, état d’une offre et assistance. Chaque notification doit conduire à une action et respecter les préférences de l’utilisateur."),
        ("Matrice fonctionnelle", "La matrice distingue présence observée, capacité annoncée officiellement et opportunité. Elle ne suppose pas l’absence d’une fonction non visible publiquement."),
        ("Matrice UX/UI", "Notation heuristique de 1 à 5 : 1 = faible, 3 = adéquat, 5 = référence. Les scores expriment une évaluation experte des parcours publics observables; ils ne mesurent ni conversion, ni satisfaction réelle, ni conformité certifiée."),
        ("Matrice CX", "La CX couvre la continuité de bout en bout : inspiration, compréhension, décision, contact, vente, transaction, suivi et soutien. Les plateformes les plus fortes réduisent l’incertitude et rendent la prochaine étape visible."),
        ("Accessibilité, responsive et contenu", "WCAG 2.2 fournit le seuil cible : contraste, reflow, navigation clavier, focus visible et non masqué, taille minimale des cibles, messages d’état et authentification accessible. La présente revue visuelle ne certifie aucun concurrent; AutoCommerce doit intégrer ces exigences dès le design system."),
        ("Scorecard et enseignements", "Cars.ca domine la simplicité et la lisibilité; CarGurus, l’intelligence décisionnelle; AutoHebdo, la couverture; Autowini, l’orchestration transfrontalière. Aucune expérience publique observée ne réunit simultanément simplicité, transparence, productivité vendeur et transaction guidée."),
        ("Espace stratégique AutoCommerce", "Positionnement recommandé : « le cockpit automobile simple et intelligent ». L’acheteur comprend rapidement le véhicule et son coût; le commerçant publie plus vite, traite mieux ses prospects et garde le contrôle des automatisations."),
        ("Blueprint de l’expérience cible", "Le service relie quatre couches : découverte, confiance, transaction et opérations. Les données VIN, médias, prix, disponibilité et interactions alimentent une chronologie partagée sans remplacer la validation humaine."),
        ("Priorisation MVP", "Le MVP doit prouver trois paris : recherche sans friction, annonce assistée de haute qualité et tableau de bord commerçant actionnable. Les fonctions complexes d’enchères, de paiement ou de logistique sont préparées architecturalement, puis activées après validation."),
        ("Recommandations finales", "S’inspirer des capacités, pas des écrans. Préserver le logo et les couleurs AutoCommerce, éclaircir les surfaces, rendre les filtres rétractables, expliciter chaque score, mesurer les parcours et conserver l’utilisateur aux commandes de l’IA."),
        ("Sources et traçabilité", "Recherche effectuée en juillet 2026. Les captures et constats peuvent évoluer. Les liens officiels et standards ci-dessous permettent de vérifier les observations et de réévaluer les scores avant une décision produit."),
    ],
}

EN = {
    "lang": "EN",
    "title": "Competitive Market Study\nUX/UI & CX Benchmark",
    "subtitle": "AUTOCOMMERCE · DOCUMENT 02 · JULY 2026",
    "toc": "Table of contents",
    "conf": "AUTOCOMMERCE · CONFIDENTIAL · EN",
    "sections": [
        ("Executive summary", "Cars.ca offers the cleanest discovery foundation; CarGurus turns pricing and reputation data into decision support; AutoTrader combines very broad inventory with a denser journey; Autowini orchestrates international transaction, inspection and shipping. AutoCommerce should not stack these models. It should unite their value in a simpler Canadian experience that merchants can operate efficiently."),
        ("Mandate, scope and method", "Expert benchmark of the public Buyer and Seller journeys of Cars.ca, CarGurus Canada, AutoTrader and Autowini. The evaluation combines interface observation, official information and UX/CX criteria. It is neither user testing, a technical audit nor a WCAG certification."),
        ("Competitive landscape", "The competitors occupy four distinct territories: discovery simplicity, market intelligence, marketplace power and cross-border commerce. This separation creates space for a platform that lowers buyer effort while increasing merchant productivity."),
        ("Cars.ca — profile", "AutoCommerce’s priority functional reference: clean interface, immediate search, Canadian bilingualism, favourites, price drops, model research and readable cards. Observed results provide compact filters, grid/list views, price badges and free CARFAX reports. The experience is efficient, while transaction and merchant tools are less visible publicly."),
        ("CarGurus — profile", "The proposition centres on trust and decision support: Deal Ratings, vehicle details, dealer reviews, budget estimate, online financing start, appointments and seller valuation. This intelligence differentiates the product, but requires transparent explanations so its scores are not perceived as a black box."),
        ("AutoTrader — profile", "AutoTrader stands out through inventory breadth, deep filtering, alerts, favourites and two selling paths: dealer offer or free private listing. This power supports expert shoppers, yet density, long descriptions, colour discontinuities and interruptions can compete with the main task."),
        ("Autowini — profile", "Autowini treats buying as an international supply chain: booking, payment, order tracking, port quality control, shipping, country guides and B2B membership. Operational value is high, but the promotional homepage, dense navigation and limited access states increase cognitive load."),
        ("Homepage visual comparison", "The four captures reveal opposite strategies: Cars.ca and CarGurus focus attention on search; AutoTrader places more navigation, promotion and consent around the action; Autowini exposes categories, offers, logistics services and membership simultaneously."),
        ("Search and discovery", "The strongest experience combines a simple entry — make, model, location — with alternate paths by budget, body type, use or recommendation. AutoCommerce should keep search central, reveal advanced criteria progressively and remember searches without forcing an early account."),
        ("Results, filters and comparison", "Cars.ca provides the best visual starting point. AutoTrader demonstrates the value of a large filter vocabulary and the risk of a dominant panel. AutoCommerce should use a collapsible drawer, removable applied-filter chips, persistent comparison and sorting that explains its effect."),
        ("Vehicle detail and trust", "An automotive decision requires total price, history, condition, photo quality, seller, availability and next steps. CarGurus value/reviews, Cars.ca readability/CARFAX and Autowini inspection/logistics should converge in a factual, explainable AutoCommerce detail page."),
        ("Seller and merchant experience", "AutoTrader clarifies the choice between dealer speed and potential private-sale value; CarGurus offers valuation and dealer connection; Autowini supports repeat professional buyers. AutoCommerce can win with a merchant cockpit: inventory, quality score, leads, offers, appointments, tasks, team and performance."),
        ("Account, retention and service", "Favourites, price alerts and saved searches support return visits. The decisive advantage is continuity: recent activity, messages, appointments, documents, offer status and help. Every notification should lead to an action and respect user preferences."),
        ("Functional matrix", "The matrix separates observed presence, official claims and opportunity. It does not infer that a feature is absent simply because it is not publicly visible."),
        ("UX/UI matrix", "Heuristic score from 1 to 5: 1 = weak, 3 = adequate, 5 = benchmark. Scores are expert evaluations of publicly observable journeys; they do not measure conversion, real satisfaction or certified compliance."),
        ("CX matrix", "CX covers end-to-end continuity: inspiration, understanding, decision, contact, selling, transaction, follow-up and support. The strongest platforms reduce uncertainty and make the next step visible."),
        ("Accessibility, responsive and content", "WCAG 2.2 defines the target threshold: contrast, reflow, keyboard access, visible and unobscured focus, minimum target size, status messages and accessible authentication. This visual review certifies no competitor; AutoCommerce should embed these requirements in its design system."),
        ("Scorecard and lessons", "Cars.ca leads simplicity and readability; CarGurus leads decision intelligence; AutoTrader leads coverage; Autowini leads cross-border orchestration. No observed public experience simultaneously unites simplicity, transparency, seller productivity and guided transaction."),
        ("AutoCommerce strategic space", "Recommended position: “the simple, intelligent automotive cockpit.” Buyers quickly understand the vehicle and its cost; merchants publish faster, manage leads better and remain in control of automation."),
        ("Target experience blueprint", "The service connects four layers: discovery, trust, transaction and operations. VIN, media, price, availability and interaction data feed a shared timeline without replacing human validation."),
        ("MVP prioritization", "The MVP must prove three bets: frictionless search, high-quality assisted listings and an actionable merchant dashboard. Complex auction, payment and logistics capabilities are prepared architecturally, then activated after validation."),
        ("Final recommendations", "Borrow capabilities, not screens. Preserve the AutoCommerce logo and colours, lighten surfaces, make filters collapsible, explain every score, measure journeys and keep users in control of AI."),
        ("Sources and traceability", "Research conducted in July 2026. Screens and findings may change. The official links and standards below support verification and future score reassessment."),
    ],
}

FUNCTION_ROWS_FR = [
    ["Capacité", "Cars.ca", "CarGurus", "AutoHebdo", "Autowini"],
    ["Recherche simple", "● Forte", "● Forte", "● Forte", "◐ Dense"],
    ["Filtres avancés", "● Compacts", "● Oui", "● Très riches", "● Par catégorie"],
    ["Favoris / alertes", "● Oui", "● Oui", "● Oui", "◐ Compte"],
    ["Aide au prix", "● Badges", "● Deal Rating", "● Outils prix", "◐ Comparaison"],
    ["Historique / inspection", "● CARFAX", "◐ Détails", "◐ Rapport", "● Inspection + VIN"],
    ["Vente particulier", "○ Peu visible", "● Estimation/dealers", "● Annonce gratuite", "● Place an AD"],
    ["Financement en ligne", "○ Non vérifié", "● Amorçage", "◐ Présent selon offre", "○ Non central"],
    ["Transaction / logistique", "○ Non vérifié", "◐ Préparation", "○ Non gérée", "● Paiement/expédition"],
    ["Cockpit B2B", "○ Non vérifié", "◐ Dealer tools", "● Marchands", "● Membership/suivi"],
]
FUNCTION_ROWS_EN = [
    ["Capability", "Cars.ca", "CarGurus", "AutoTrader", "Autowini"],
    ["Simple search", "● Strong", "● Strong", "● Strong", "◐ Dense"],
    ["Advanced filters", "● Compact", "● Yes", "● Very deep", "● Category-led"],
    ["Favourites / alerts", "● Yes", "● Yes", "● Yes", "◐ Account"],
    ["Price support", "● Badges", "● Deal Rating", "● Pricing tools", "◐ Comparison"],
    ["History / inspection", "● CARFAX", "◐ Details", "◐ Report", "● Inspection + VIN"],
    ["Private selling", "○ Less visible", "● Valuation/dealers", "● Free listing", "● Place an AD"],
    ["Online financing", "○ Not verified", "● Start online", "◐ Offer-dependent", "○ Not central"],
    ["Transaction / logistics", "○ Not verified", "◐ Preparation", "○ Not managed", "● Pay/ship"],
    ["B2B cockpit", "○ Not verified", "◐ Dealer tools", "● Merchants", "● Membership/tracking"],
]

UX_ROWS = [
    ("Clarity / hierarchy", 5, 4, 3, 2),
    ("Navigation / IA", 4, 4, 3, 2),
    ("Search entry", 5, 5, 4, 3),
    ("Filters / results", 5, 4, 3, 3),
    ("Visual consistency", 5, 4, 3, 2),
    ("Decision support", 4, 5, 4, 4),
    ("Responsive confidence*", 4, 4, 3, 3),
    ("Accessibility readiness*", 4, 4, 3, 2),
]
CX_ROWS = [
    ("Discovery", 5, 5, 4, 3),
    ("Price transparency", 4, 5, 4, 4),
    ("Trust evidence", 4, 5, 4, 5),
    ("Seller enablement", 3, 5, 5, 4),
    ("Transaction continuity", 2, 4, 3, 5),
    ("Retention / alerts", 4, 5, 5, 4),
    ("Support visibility", 3, 5, 4, 4),
    ("Cross-border readiness", 1, 1, 1, 5),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def style_document(doc: Document, conf: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    for name, size, color in [("Title", 32, BLACK), ("Heading 1", 22, BLACK), ("Heading 2", 13, RED)]:
        st = doc.styles[name]
        st.font.name = "Aptos Display"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(3)
        st.paragraph_format.space_after = Pt(8)
    header = section.header.paragraphs[0]
    header.text = conf
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(7)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(GREY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("AUTOCOMMERCE  ·  ")
    add_field(footer, "PAGE")
    for run in footer.runs:
        run.font.size = Pt(7)
        run.font.color.rgb = RGBColor.from_string(GREY)


def add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(RED)


def add_title(doc: Document, number: int, title: str, summary: str) -> None:
    add_kicker(doc, f"{number:02d} · Competitive benchmark")
    doc.add_heading(title, 0)
    p = doc.add_paragraph(summary)
    p.paragraph_format.space_after = Pt(10)
    p.runs[0].font.size = Pt(10.2)
    p.runs[0].font.color.rgb = RGBColor.from_string(GREY)


def add_callout(doc: Document, title: str, body: str, color: str = GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(1)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc: Document, rows: list[list[str]], widths: list[float] | None = None, font_size: float = 7.4) -> None:
    table = doc.add_table(rows=0, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, values in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(values):
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[j].width = Inches(widths[j])
            p = cells[j].paragraphs[0]
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
            r = p.add_run(str(value))
            r.font.size = Pt(font_size)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor.from_string(WHITE)
                set_cell_shading(cells[j], BLACK if j == 0 else GREEN)
            elif j == 0:
                r.bold = True
                set_cell_shading(cells[j], LIGHT)
            elif "●" in str(value):
                r.font.color.rgb = RGBColor.from_string(GREEN)
            elif "○" in str(value):
                r.font.color.rgb = RGBColor.from_string(RED)
        if i == 0:
            set_repeat_table_header(table.rows[-1])
    doc.add_paragraph()


def add_score_table(doc: Document, rows, language: str, note: str) -> None:
    head = ["Critère" if language == "FR" else "Criterion", "Cars.ca", "CarGurus", "AutoTrader", "Autowini"]
    data = [head]
    for label, *scores in rows:
        data.append([label] + [f"{s}/5  {'●' * s}{'○' * (5-s)}" for s in scores])
    avg = [sum(r[i] for r in rows) / len(rows) for i in range(1, 5)]
    data.append(["MOYENNE" if language == "FR" else "AVERAGE"] + [f"{x:.1f}/5" for x in avg])
    add_table(doc, data, [2.05, 1.23, 1.23, 1.23, 1.23], 7.0)
    add_callout(doc, "Lecture" if language == "FR" else "Reading note", note, BLUE)


def add_capture(doc: Document, path: Path, caption: str, width: float = 6.95) -> None:
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        c = doc.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(5)
        c.runs[0].italic = True
        c.runs[0].font.size = Pt(7.2)
        c.runs[0].font.color.rgb = RGBColor.from_string(GREY)


def add_four_layer_diagram(doc: Document, language: str) -> None:
    labels = (
        ["1 · DÉCOUVERTE", "2 · CONFIANCE", "3 · TRANSACTION", "4 · OPÉRATIONS"]
        if language == "FR"
        else ["1 · DISCOVERY", "2 · TRUST", "3 · TRANSACTION", "4 · OPERATIONS"]
    )
    bodies = (
        ["Recherche · filtres rétractables · comparaison", "VIN · photos · historique · prix explicable", "Offre · rendez-vous · documents · chronologie", "Inventaire · prospects · tâches · analytique"]
        if language == "FR"
        else ["Search · collapsible filters · comparison", "VIN · photos · history · explainable price", "Offer · appointment · documents · timeline", "Inventory · leads · tasks · analytics"]
    )
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    colors = [BLACK, GREEN, BLUE, RED]
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, colors[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(labels[i])
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(WHITE)
        p2 = cell.add_paragraph(bodies[i])
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for rr in p2.runs:
            rr.font.size = Pt(7.2)
            rr.font.color.rgb = RGBColor.from_string(WHITE)
    doc.add_paragraph("→            →            →", style=None).alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_specific(doc: Document, idx: int, lang: str) -> None:
    fr = lang == "FR"
    if idx == 1:
        add_callout(doc, "Décision" if fr else "Decision", "Construire le meilleur socle canadien de découverte et de vente assistée — sans reproduire l’interface d’un concurrent." if fr else "Build the strongest Canadian foundation for discovery and assisted selling—without reproducing a competitor’s interface.")
        add_table(doc, [
            ["Référence" if fr else "Reference", "Valeur à retenir" if fr else "Value to retain", "Limite à dépasser" if fr else "Limit to exceed"],
            ["Cars.ca", "Clarté, filtres compacts, CARFAX", "Peu de continuité transactionnelle publique" if fr else "Limited public transaction continuity"],
            ["CarGurus", "Intelligence prix, avis, financement", "Rendre les scores explicables" if fr else "Make scores explainable"],
            ["AutoTrader", "Couverture, vente, filtres", "Réduire densité et interruptions" if fr else "Reduce density and interruptions"],
            ["Autowini", "Inspection, paiement, expédition", "Simplifier IA et surfaces" if fr else "Simplify IA and surfaces"],
        ], [1.25, 2.7, 3.0])
    elif idx == 2:
        add_table(doc, [
            ["Dimension", "Protocole" if fr else "Protocol"],
            ["Evidence", "Capture datée + observation DOM + page officielle" if fr else "Dated capture + DOM observation + official page"],
            ["UX/UI", "Heuristiques, hiérarchie, navigation, recherche, filtres, cohérence, responsive"],
            ["CX", "Effort, confiance, transparence, continuité, soutien, fidélisation"],
            ["Score", "1 faible · 3 adéquat · 5 référence" if fr else "1 weak · 3 adequate · 5 benchmark"],
            ["Limite", "Pas de données internes, analytics ou tests utilisateurs" if fr else "No internal data, analytics or user tests"],
        ], [1.35, 5.6])
    elif idx == 3:
        add_table(doc, [
            ["Plateforme" if fr else "Platform", "Territoire" if fr else "Territory", "Acheteur" if fr else "Buyer", "Vendeur / B2B" if fr else "Seller / B2B"],
            ["Cars.ca", "Découverte simple" if fr else "Simple discovery", "Trouver vite" if fr else "Find quickly", "Visibilité publique limitée" if fr else "Limited public visibility"],
            ["CarGurus", "Intelligence marché" if fr else "Market intelligence", "Évaluer l’offre" if fr else "Assess the deal", "Valeur + leads" if fr else "Valuation + leads"],
            ["AutoTrader", "Place de marché" if fr else "Marketplace scale", "Explorer en profondeur" if fr else "Explore deeply", "Dealer ou privé" if fr else "Dealer or private"],
            ["Autowini", "Commerce mondial" if fr else "Global trade", "Importer avec suivi" if fr else "Import with tracking", "Membership volume"],
        ], [1.25, 2.0, 1.85, 1.85])
        add_callout(doc, "White space", "Simplicité de Cars.ca + confiance de CarGurus + opérations d’Autowini, dans une interface plus légère qu’AutoTrader." if fr else "Cars.ca simplicity + CarGurus trust + Autowini operations, in an interface lighter than AutoTrader.", BLUE)
    elif idx in (4, 5, 6, 7):
        names = ["Cars.ca", "CarGurus", "AutoTrader", "Autowini"]
        keys = ["cars", "cargurus", "autotrader", "autowini"]
        strengths = [
            ["Hiérarchie nette", "Recherche immédiate", "Filtres non dominants", "Cartes lisibles"],
            ["Deal Ratings", "Avis concessionnaires", "Budget / financement", "Estimation vendeur"],
            ["Inventaire profond", "Filtres exhaustifs", "Alertes / favoris", "Deux voies de vente"],
            ["Prix rendu / shipping", "Inspection + QC", "Suivi commande", "Membership B2B"],
        ]
        risks = [
            "Rendre visibles la transaction et le cockpit commerçant." if fr else "Make transaction and merchant cockpit more visible.",
            "Expliquer les modèles de score et la provenance des données." if fr else "Explain scoring models and data provenance.",
            "Réduire la densité et préserver la tâche face aux interruptions." if fr else "Reduce density and protect the task from interruptions.",
            "Réduire la charge cognitive et moderniser la cohérence visuelle." if fr else "Lower cognitive load and modernize visual consistency.",
        ]
        pos = idx - 4
        add_capture(doc, CAPTURES[keys[pos]], f"Capture de référence · {names[pos]} · juillet 2026" if fr else f"Reference capture · {names[pos]} · July 2026", 6.7)
        add_bullets(doc, strengths[pos])
        add_callout(doc, "Point de vigilance" if fr else "Watchpoint", risks[pos], AMBER)
    elif idx == 8:
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, key, cap in zip([c for r in table.rows for c in r.cells], ["cars", "cargurus", "autotrader", "autowini"], ["Cars.ca", "CarGurus", "AutoTrader", "Autowini"]):
            if CAPTURES[key].exists():
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(CAPTURES[key]), width=Inches(3.2))
                q = cell.add_paragraph(cap)
                q.alignment = WD_ALIGN_PARAGRAPH.CENTER
                q.runs[0].bold = True
                q.runs[0].font.size = Pt(7.5)
        add_callout(doc, "Lecture" if fr else "Reading", "La quantité d’éléments visibles avant interaction est un choix stratégique : elle influence la vitesse de compréhension, pas seulement l’esthétique." if fr else "The number of elements visible before interaction is strategic: it affects comprehension speed, not only aesthetics.")
    elif idx == 9:
        add_table(doc, [
            ["Critère" if fr else "Criterion", "Cars.ca", "CarGurus", "AutoTrader", "Autowini", "Cible AutoCommerce" if fr else "AutoCommerce target"],
            ["Entrée", "Marque/modèle/lieu", "Marque/modèle/CP", "Marque/modèle/CP", "Recherche + tags", "Question courte + alternatives"],
            ["Exploration", "Recommandations", "Catégorie/budget", "Type/marque", "Catégorie/pays", "Budget, usage, véhicule"],
            ["Progressivité", "Forte", "Forte", "Moyenne", "Faible", "Forte, sans blocage"],
            ["Mémoire", "Recommandations", "Recherches récentes", "Favoris/alertes", "Compte", "Locale puis compte"],
        ], [1.25, 1.18, 1.25, 1.25, 1.15, 1.35], 6.8)
    elif idx == 10:
        add_capture(doc, CAPTURES["cars_results"], "Cars.ca · résultats observés : filtres compacts, vues et badges." if fr else "Cars.ca · observed results: compact filters, views and badges.", 6.55)
        add_table(doc, [
            ["Principe cible" if fr else "Target principle", "Comportement" if fr else "Behaviour", "Bénéfice" if fr else "Benefit"],
            ["Tiroir rétractable" if fr else "Collapsible drawer", "Fermé après application, état conservé" if fr else "Closes after apply, state retained", "Inventaire dominant" if fr else "Inventory stays dominant"],
            ["Pastilles actives" if fr else "Applied chips", "Retrait individuel / tout effacer" if fr else "Remove one / clear all", "Contrôle visible" if fr else "Visible control"],
            ["Comparer", "2–4 véhicules persistants" if fr else "2–4 persistent vehicles", "Décision plus sûre" if fr else "Safer decision"],
        ], [2.0, 3.0, 2.0])
    elif idx == 11:
        add_table(doc, [
            ["Bloc" if fr else "Block", "Preuve requise" if fr else "Required evidence", "Traitement AutoCommerce" if fr else "AutoCommerce treatment"],
            ["Prix", "Total + frais + financement indicatif" if fr else "Total + fees + indicative financing", "Aucun prix conditionnel caché" if fr else "No hidden conditional price"],
            ["État", "Photos guidées + dommages + inspection" if fr else "Guided photos + damage + inspection", "Faits validés, horodatés" if fr else "Validated, timestamped facts"],
            ["Historique", "VIN + rapport + provenance" if fr else "VIN + report + provenance", "Source et date visibles" if fr else "Visible source and date"],
            ["Vendeur", "Identité + réponse + avis" if fr else "Identity + response + reviews", "Score explicable" if fr else "Explainable score"],
            ["Action", "Essai, message, offre, achat" if fr else "Test drive, message, offer, buy", "Une CTA principale contextuelle" if fr else "One contextual primary CTA"],
        ], [1.2, 2.8, 3.0])
        add_callout(doc, "Règle", "Une IA peut suggérer une description; elle ne doit jamais inventer l’état, l’historique ou la disponibilité." if fr else "AI may suggest a description; it must never invent condition, history or availability.", RED)
    elif idx == 12:
        add_table(doc, [
            ["Moment" if fr else "Moment", "Friction actuelle du marché" if fr else "Market friction", "Réponse AutoCommerce" if fr else "AutoCommerce response"],
            ["Onboarding", "Saisie répétitive", "VIN + import + sauvegarde automatique" if fr else "VIN + import + autosave"],
            ["Photos", "Qualité variable", "Guide d’angles + contrôle flou/doublons" if fr else "Shot guide + blur/duplicate check"],
            ["Publication", "Formulaire long", "Progressif + score de complétude" if fr else "Progressive + completion score"],
            ["Prospects", "Canaux dispersés", "Inbox, SLA, tâches, attribution" if fr else "Inbox, SLA, tasks, assignment"],
            ["Pilotage", "Données sans action", "Insights liés à une recommandation" if fr else "Insights tied to a recommendation"],
        ], [1.2, 2.7, 3.1])
    elif idx == 13:
        add_four_layer_diagram(doc, lang)
        add_bullets(doc, [
            "Acheteur : recherches, favoris, comparaison, messages, rendez-vous, offres, documents." if fr else "Buyer: searches, favourites, comparison, messages, appointments, offers, documents.",
            "Commerçant : inventaire, prospects, tâches, équipe, performance et facturation." if fr else "Merchant: inventory, leads, tasks, team, performance and billing.",
            "AutoCommerce : qualité, modération, fraude, litiges, abonnements et audit." if fr else "AutoCommerce: quality, moderation, fraud, disputes, subscriptions and audit.",
        ])
    elif idx == 14:
        add_table(doc, FUNCTION_ROWS_FR if fr else FUNCTION_ROWS_EN, [1.65, 1.35, 1.35, 1.35, 1.35], 6.8)
        add_callout(doc, "Légende" if fr else "Legend", "● observé/confirmé · ◐ partiel ou dépendant du contexte · ○ non vérifié publiquement" if fr else "● observed/confirmed · ◐ partial or context-dependent · ○ not publicly verified", BLUE)
    elif idx == 15:
        add_score_table(doc, UX_ROWS, lang, "Les deux critères marqués * sont des indices visuels, pas des tests de conformité." if fr else "The two * criteria are visual indicators, not conformance tests.")
        add_bullets(doc, [
            "Cars.ca : référence de hiérarchie et de progressivité." if fr else "Cars.ca: hierarchy and progressive-disclosure reference.",
            "CarGurus : meilleure intégration de la décision dans l’interface." if fr else "CarGurus: strongest decision support embedded in the interface.",
            "AutoTrader : puissance fonctionnelle à simplifier." if fr else "AutoTrader: functional power that needs simplification.",
            "Autowini : profondeur de service à restructurer." if fr else "Autowini: service depth that needs restructuring.",
        ])
    elif idx == 16:
        add_score_table(doc, CX_ROWS, lang, "Les scores portent sur la continuité publique observée et les capacités officielles, pas sur le NPS ou la qualité réelle du support." if fr else "Scores reflect observed public continuity and official capabilities, not NPS or actual support quality.")
    elif idx == 17:
        add_table(doc, [
            ["Exigence WCAG 2.2" if fr else "WCAG 2.2 requirement", "Risque automobile" if fr else "Automotive risk", "Critère MVP" if fr else "MVP criterion"],
            ["1.4.3 / 1.4.11 Contraste", "Badges et états illisibles" if fr else "Unreadable badges and states", "Texte, composants, graphiques testés" if fr else "Test text, components, charts"],
            ["1.4.10 Reflow", "Filtres hors écran" if fr else "Filters off-screen", "320 CSS px sans perte" if fr else "320 CSS px without loss"],
            ["2.4.7 / 2.4.11 Focus", "Navigation clavier invisible" if fr else "Invisible keyboard navigation", "Focus visible et non masqué" if fr else "Visible, unobscured focus"],
            ["2.5.8 Cibles", "Icônes mobiles difficiles" if fr else "Hard mobile icons", "24×24 CSS px minimum"],
            ["4.1.3 Statuts", "Résultats / sauvegarde silencieux" if fr else "Silent results / save", "Messages annoncés aux aides" if fr else "Assistive status messages"],
        ], [1.85, 2.35, 2.95], 7.1)
        add_callout(doc, "Responsive", "Une même priorité, des compositions adaptées : tiroir desktop, bottom sheet mobile, actions au pouce et aucun filtre dominant." if fr else "Same priority, adapted compositions: desktop drawer, mobile bottom sheet, thumb-friendly actions and no dominant filter.")
    elif idx == 18:
        add_table(doc, [
            ["Plateforme", "Force signature" if fr else "Signature strength", "Risque principal" if fr else "Primary risk", "Leçon AutoCommerce" if fr else "AutoCommerce lesson"],
            ["Cars.ca", "Simplicité 5/5", "Profondeur moins visible" if fr else "Less visible depth", "Socle UX"],
            ["CarGurus", "Décision 5/5", "Opacité algorithmique" if fr else "Algorithmic opacity", "Confiance explicable" if fr else "Explainable trust"],
            ["AutoTrader", "Couverture 5/5", "Densité", "Progressive disclosure"],
            ["Autowini", "Opérations 5/5", "Complexité", "Timeline transactionnelle" if fr else "Transaction timeline"],
        ], [1.25, 1.7, 2.05, 2.25])
    elif idx == 19:
        add_four_layer_diagram(doc, lang)
        add_table(doc, [
            ["Différenciateur" if fr else "Differentiator", "Promesse" if fr else "Promise", "Preuve produit" if fr else "Product proof"],
            ["Filtres adaptatifs" if fr else "Adaptive filters", "Voir plus de véhicules, moins d’interface" if fr else "See more vehicles, less interface", "Tiroir + chips + vues"],
            ["Annonce intelligente" if fr else "Intelligent listing", "Publier juste du premier coup" if fr else "Publish right the first time", "VIN + photo QA + texte validé"],
            ["Confiance explicable" if fr else "Explainable trust", "Comprendre chaque score" if fr else "Understand every score", "Source, date, facteurs"],
            ["Cockpit commerçant" if fr else "Merchant cockpit", "Savoir quoi faire ensuite" if fr else "Know what to do next", "SLA, tâches, alertes, performance"],
        ], [1.55, 2.75, 2.7])
    elif idx == 20:
        add_four_layer_diagram(doc, lang)
        add_table(doc, [
            ["Acteur" if fr else "Actor", "Action", "Système" if fr else "System", "Preuve / état" if fr else "Evidence / state"],
            ["Acheteur" if fr else "Buyer", "Recherche → compare → contacte" if fr else "Search → compare → contact", "Search + recommendation", "Filtres et favoris" if fr else "Filters and favourites"],
            ["Commerçant" if fr else "Merchant", "Importe → valide → publie" if fr else "Import → validate → publish", "VIN/media/listing", "Score qualité" if fr else "Quality score"],
            ["Équipe" if fr else "Team", "Répond → planifie → conclut" if fr else "Respond → schedule → close", "CRM workflow", "SLA + timeline"],
            ["AutoCommerce", "Modère → mesure → améliore" if fr else "Moderate → measure → improve", "Rules + analytics", "Audit + KPI"],
        ], [1.25, 2.25, 1.8, 1.7], 7.0)
    elif idx == 21:
        add_table(doc, [
            ["Priorité" if fr else "Priority", "Capacité" if fr else "Capability", "Valeur" if fr else "Value", "Décision prototype" if fr else "Prototype decision"],
            ["P0", "Recherche + résultats + filtre rétractable" if fr else "Search + results + collapsible filter", "Prouver la simplicité" if fr else "Prove simplicity", "Test desktop/mobile"],
            ["P0", "VIN + photos + description assistée" if fr else "VIN + photos + assisted description", "Réduire temps/erreurs" if fr else "Lower time/errors", "Validation humaine obligatoire" if fr else "Human validation required"],
            ["P0", "Cockpit inventaire/prospects", "Actionnabilité", "États + SLA + tâches"],
            ["P1", "Favoris, alertes, comparaison", "Rétention", "Compte progressif"],
            ["P1", "Prix/confiance explicables" if fr else "Explainable price/trust", "Conversion", "Facteurs et sources"],
            ["P2", "Offres, enchères, paiement, shipping", "Revenus", "Après validation P0/P1" if fr else "After P0/P1 validation"],
        ], [0.75, 2.65, 1.65, 2.0], 7.1)
    elif idx == 22:
        add_table(doc, [
            ["Horizon", "Action", "Mesure de succès" if fr else "Success measure"],
            ["0–30 j", "Prototype Figma des trois paris P0" if fr else "Figma prototype of the three P0 bets", "5 scénarios testables" if fr else "5 testable scenarios"],
            ["31–60 j", "Tests utilisateurs acheteurs + commerçants" if fr else "Buyer + merchant user testing", "Taux de réussite, temps, erreurs" if fr else "Success, time, errors"],
            ["61–90 j", "Scenario Make contrôlé + instrumentation" if fr else "Controlled Make scenario + instrumentation", "Publication assistée traçable" if fr else "Traceable assisted publishing"],
            ["Après MVP" if fr else "Post-MVP", "Transaction, enchères, logistique" if fr else "Transaction, auction, logistics", "Preuve de demande et conformité" if fr else "Demand and compliance evidence"],
        ], [1.05, 3.75, 2.25])
        add_callout(doc, "North Star", "Temps médian entre l’intention et la prochaine action utile — séparé pour acheteurs et commerçants." if fr else "Median time from intent to the next useful action—tracked separately for buyers and merchants.", GREEN)
    elif idx == 23:
        rows = [["#", "Source", "URL"]]
        for i, (name, url) in enumerate(SOURCES, 1):
            rows.append([str(i), name, url])
        add_table(doc, rows, [0.35, 3.1, 3.65], 5.9)
        add_callout(doc, "Traçabilité" if fr else "Traceability", "Les évaluations doivent être révisées avant le lancement si les interfaces, conditions ou fonctions concurrentes changent." if fr else "Reassess findings before launch if competitor interfaces, terms or capabilities change.", BLUE)


def cover(doc: Document, data: dict) -> None:
    doc.add_paragraph("\n\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AUTOCOMMERCE")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(RED)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(50)
    title.paragraph_format.space_after = Pt(20)
    rr = title.add_run(data["title"])
    rr.bold = True
    rr.font.name = "Aptos Display"
    rr.font.size = Pt(31)
    rr.font.color.rgb = RGBColor.from_string(BLACK)
    line = doc.add_paragraph("━" * 36)
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.runs[0].font.color.rgb = RGBColor.from_string(RED)
    sub = doc.add_paragraph(data["subtitle"])
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor.from_string(GREY)
    doc.add_paragraph("\n\n")
    for i, text in enumerate(SIGNATURE):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = i < 2
        r.font.size = Pt(9 if i != 1 else 13)
        r.font.color.rgb = RGBColor.from_string(BLACK if i else GREEN)


def toc(doc: Document, data: dict) -> None:
    add_kicker(doc, "AUTOCommerce · Document 02")
    doc.add_heading(data["toc"], 0)
    rows = []
    for i, (title, _) in enumerate(data["sections"], 1):
        rows.append([f"{i:02d}", title, str(i + 2)])
    add_table(doc, [["", "Section", "Page"]] + rows, [0.45, 5.8, 0.65], 7.2)


def generate(data: dict) -> Path:
    doc = Document()
    style_document(doc, data["conf"])
    cover(doc, data)
    doc.add_page_break()
    toc(doc, data)
    for idx, (title, summary) in enumerate(data["sections"], 1):
        doc.add_page_break()
        add_title(doc, idx, title, summary)
        add_page_specific(doc, idx, data["lang"])
    OUTPUT.mkdir(parents=True, exist_ok=True)
    target = OUTPUT / f"autocommerce-document-02-{data['lang'].lower()}.docx"
    doc.save(target)
    return target


if __name__ == "__main__":
    for payload in (FR, EN):
        print(generate(payload))
