from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = "14213D"
MUTED = "64748B"
ACCENT = "006D77"
ACCENT_SOFT = "E8F3F3"
LINE = "DBE4EA"
SURFACE = "F6F8FA"


class WordDocumentRenderer:
    """Create editable executive DOCX reports from ATELIER Markdown."""

    def render(self, *, title: str, content: str, output_path: Path) -> Path:
        document = Document()
        self._configure_document(document)
        self._add_cover(document, title)
        document.add_page_break()
        self._add_toc(document, content)
        document.add_page_break()
        self._add_markdown(document, content)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        return output_path

    def _configure_document(self, document: Document) -> None:
        section = document.sections[0]
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string(INK)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1
        self._set_style_font(normal, "Arial")

        for name, size, color, before, after in (
            ("Heading 1", 20, INK, 16, 8),
            ("Heading 2", 15, INK, 12, 6),
            ("Heading 3", 12, ACCENT, 8, 4),
        ):
            style = document.styles[name]
            style.font.name = "Arial"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True
            self._set_style_font(style, "Arial")

        footer = section.footer
        table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        left, right = table.rows[0].cells
        left.text = "NEXUS CORE v2 · CONFIDENTIEL"
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        right.paragraphs[0].add_run("PAGE ")
        self._add_field(right.paragraphs[0], "PAGE")
        for cell in (left, right):
            for run in cell.paragraphs[0].runs:
                self._format_run(run, size=8, color=MUTED)

    def _add_cover(self, document: Document, title: str) -> None:
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(125)
        eyebrow = document.add_paragraph()
        eyebrow.paragraph_format.space_after = Pt(22)
        self._format_run(
            eyebrow.add_run("NEXUS CORE v2 · EXECUTIVE REPORT"),
            size=9,
            color=ACCENT,
            bold=True,
        )
        title_paragraph = document.add_paragraph()
        title_paragraph.paragraph_format.space_after = Pt(22)
        title_paragraph.paragraph_format.keep_with_next = True
        self._format_run(
            title_paragraph.add_run(title),
            size=30,
            color=INK,
            bold=True,
        )
        rule = document.add_paragraph()
        rule.paragraph_format.space_after = Pt(20)
        self._paragraph_bottom_border(rule, ACCENT, size=18)

        prepared = document.add_paragraph()
        prepared.paragraph_format.space_after = Pt(2)
        self._format_run(prepared.add_run("PREPARED BY"), size=8, color=MUTED)
        name = document.add_paragraph()
        name.paragraph_format.space_after = Pt(2)
        self._format_run(name.add_run("DJIGO DJIBI"), size=13, color=INK, bold=True)
        role = document.add_paragraph()
        self._format_run(
            role.add_run(
                "CX Consultant | Strategic Product & UX/UI Designer | AI Front-End Developer"
            ),
            size=8,
            color=MUTED,
        )

    def _add_toc(self, document: Document, content: str) -> None:
        eyebrow = document.add_paragraph()
        self._format_run(eyebrow.add_run("SOMMAIRE"), size=9, color=ACCENT, bold=True)
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(30)
        heading.paragraph_format.space_after = Pt(24)
        self._format_run(
            heading.add_run("Table des matières"),
            size=24,
            color=INK,
            bold=True,
        )
        for match in re.finditer(r"^(#{2,3})\s+(.+?)\s*$", content, re.MULTILINE):
            level = len(match.group(1))
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.3 if level == 3 else 0)
            paragraph.paragraph_format.space_after = Pt(7)
            self._format_run(
                paragraph.add_run(match.group(2)),
                size=10 if level == 3 else 11,
                color=MUTED if level == 3 else INK,
                bold=level == 2,
            )

    def _add_markdown(self, document: Document, content: str) -> None:
        lines = content.splitlines()
        buffer: list[str] = []
        index = 0

        def flush() -> None:
            if buffer:
                paragraph = document.add_paragraph()
                self._add_inline(paragraph, " ".join(buffer))
                buffer.clear()

        while index < len(lines):
            line = lines[index].rstrip()
            if line.startswith("### "):
                flush()
                document.add_paragraph(line[4:], style="Heading 3")
            elif line.startswith("## "):
                flush()
                document.add_paragraph(line[3:], style="Heading 2")
            elif line.startswith("> "):
                flush()
                self._add_callout(document, line[2:])
            elif re.match(r"^[-*]\s+", line):
                flush()
                paragraph = document.add_paragraph(style="List Bullet")
                self._add_inline(paragraph, line[2:])
            elif re.match(r"^\d+\.\s+", line):
                flush()
                paragraph = document.add_paragraph(style="List Number")
                self._add_inline(paragraph, re.sub(r"^\d+\.\s+", "", line))
            elif line.startswith("|") and index + 1 < len(lines):
                flush()
                table_lines = [line]
                index += 1
                table_lines.append(lines[index])
                while index + 1 < len(lines) and lines[index + 1].startswith("|"):
                    index += 1
                    table_lines.append(lines[index])
                self._add_table(document, table_lines)
            elif not line.strip():
                flush()
            else:
                buffer.append(line)
            index += 1
        flush()

    def _add_callout(self, document: Document, text: str) -> None:
        table = document.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        self._set_cell_fill(cell, ACCENT_SOFT)
        self._set_cell_margins(cell, 120, 120, 180, 180)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        self._add_inline(cell.paragraphs[0], text)
        document.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_table(self, document: Document, lines: list[str]) -> None:
        rows = [[cell.strip() for cell in line.strip("|").split("|")] for line in lines]
        if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
            rows.pop(1)
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        column_width = Inches(6.5 / len(rows[0]))
        for row_index, row in enumerate(rows):
            for column_index, value in enumerate(row):
                cell = table.cell(row_index, column_index)
                cell.width = column_width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                self._set_cell_margins(cell, 80, 80, 120, 120)
                if row_index == 0:
                    self._set_cell_fill(cell, INK)
                elif row_index % 2 == 0:
                    self._set_cell_fill(cell, SURFACE)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.space_after = Pt(0)
                self._format_run(
                    paragraph.add_run(value),
                    size=8.5,
                    color="FFFFFF" if row_index == 0 else INK,
                    bold=row_index == 0,
                )
        document.add_paragraph().paragraph_format.space_after = Pt(2)

    def _add_inline(self, paragraph, text: str) -> None:
        parts = re.split(r"(\*\*.+?\*\*|`.+?`)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                self._format_run(paragraph.add_run(part[2:-2]), bold=True)
            elif part.startswith("`") and part.endswith("`"):
                run = paragraph.add_run(part[1:-1])
                self._format_run(run, name="Courier New", size=9)
            else:
                self._format_run(paragraph.add_run(part))

    @staticmethod
    def _format_run(
        run,
        *,
        name: str = "Arial",
        size: float | None = None,
        color: str | None = None,
        bold: bool | None = None,
    ) -> None:
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
        if size is not None:
            run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = RGBColor.from_string(color)
        if bold is not None:
            run.bold = bold

    @staticmethod
    def _set_style_font(style, name: str) -> None:
        style.element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
        style.element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)

    @staticmethod
    def _set_cell_fill(cell, color: str) -> None:
        properties = cell._tc.get_or_add_tcPr()
        shading = properties.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            properties.append(shading)
        shading.set(qn("w:fill"), color)

    @staticmethod
    def _set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:
        properties = cell._tc.get_or_add_tcPr()
        margins = properties.first_child_found_in("w:tcMar")
        if margins is None:
            margins = OxmlElement("w:tcMar")
            properties.append(margins)
        for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
            node = margins.find(qn(f"w:{edge}"))
            if node is None:
                node = OxmlElement(f"w:{edge}")
                margins.append(node)
            node.set(qn("w:w"), str(value))
            node.set(qn("w:type"), "dxa")

    @staticmethod
    def _paragraph_bottom_border(paragraph, color: str, size: int) -> None:
        properties = paragraph._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color)
        borders.append(bottom)
        properties.append(borders)

    @staticmethod
    def _add_field(paragraph, instruction: str) -> None:
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), instruction)
        paragraph._p.append(field)
