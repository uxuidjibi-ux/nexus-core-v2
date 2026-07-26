from nexus.document_renderer import ExecutiveDocumentRenderer


def test_renderer_creates_chartered_html(tmp_path):
    renderer = ExecutiveDocumentRenderer()
    result = renderer.render(
        title="Analyse Fonctionnelle",
        content="## Introduction\n\nTexte clair.\n\n## Décision\n\n> **POINT CLÉ —** MVP.",
        output_dir=tmp_path,
        slug="analyse",
        export_pdf=False,
        export_docx=False,
    )
    html = result.html_path.read_text(encoding="utf-8")
    assert "PREPARED BY" in html
    assert "DJIGO DJIBI" in html
    assert "Table des matières" in html
    assert 'id="introduction"' in html
    assert result.pdf_path is None
    assert result.docx_path is None


def test_renderer_creates_editable_word_document(tmp_path):
    from docx import Document

    renderer = ExecutiveDocumentRenderer()
    result = renderer.render(
        title="Analyse Fonctionnelle",
        content=(
            "## Introduction\n\nTexte clair.\n\n"
            "## Comparatif\n\n| Option | Décision |\n|---|---|\n| A | Oui |"
        ),
        output_dir=tmp_path,
        slug="analyse-word",
        export_pdf=False,
    )
    assert result.docx_path is not None
    document = Document(result.docx_path)
    full_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "PREPARED BY" in full_text
    assert "DJIGO DJIBI" in full_text
    assert "Table des matières" in full_text
    assert len(document.tables) >= 1


def test_autocommerce_has_four_document_specs():
    from agents.creative_dir import AUTOCOMMERCE_DOCUMENTS

    assert len(AUTOCOMMERCE_DOCUMENTS) == 4
    assert {item[0] for item in AUTOCOMMERCE_DOCUMENTS} == {
        "analyse-fonctionnelle",
        "benchmark-marche",
        "methodologie-appliquee",
        "specifications-mvp",
    }


def test_atelier_profile_is_registered():
    from agents import ALL_AGENTS

    assert [agent.name for agent in ALL_AGENTS][-1] == "ATELIER"
